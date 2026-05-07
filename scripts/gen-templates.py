# -*- coding: utf-8 -*-
"""
T001 / T011 / T016 のテンプレートファイル（Word + Excel）を生成する。
出力先: ../public/files/

実行方法:
  cd sites/template-free-jp/web
  python scripts/gen-templates.py
"""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter


OUT_DIR = Path(__file__).parent.parent / "public" / "files"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_a4_portrait(section, b5=False):
    """A4縦（または B5）のページサイズを設定"""
    if b5:
        section.page_width = Cm(18.2)
        section.page_height = Cm(25.7)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def add_footer_note(doc):
    """共通フッター: ユーザーがそのまま実用できるよう、テンプレ内には何も追記しない。
    免責・配布元情報はサイト本体（利用規約・各テンプレページ）に記載済み。"""
    return


# ==========================================================
# T001: 退職届
# ==========================================================
def gen_taishokutodoke_a4_tategaki():
    """A4縦書き退職届"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    # 縦書き設定
    section = doc.sections[0]
    sectPr = section._sectPr
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'tbRl')
    sectPr.append(textDirection)

    # 表題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("退　職　届")
    run.font.name = "ＭＳ 明朝"
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 本文
    body = (
        "　私事、",
        "",
        "　この度、一身上の都合により、",
        "　令和　年　月　日をもって退職いたします。",
        "",
        "",
        "　令和　年　月　日",
        "",
        "　　　　　　　　　　所属：　　　　　　　　　　　",
        "　　　　　　　　　　氏名：　　　　　　　　　　　　㊞",
        "",
        "",
        "○○株式会社",
        "代表取締役　○○　○○　殿",
    )
    for line in body:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "ＭＳ 明朝"
        run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "taishokutodoke_a4_tategaki.docx")


def gen_taishokutodoke_a4_yokogaki():
    """A4横書き退職届"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    # 表題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("退　職　届")
    run.font.name = "游明朝"
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph("")

    # 日付・宛名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    doc.add_paragraph("")

    # 所属・氏名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("所属：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # 本文
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "　私事、この度、一身上の都合により、令和　年　月　日をもって"
        "退職いたします。\n\n　以上"
    )
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    # 添え状
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("送　付　状")
    run.font.name = "游明朝"
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n人事部　御中")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("〒　　　-　　　\n住所：　　　　　　　　　　\n氏名：　　　　　　　　　　")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(
        "拝啓　時下ますますご清栄のこととお慶び申し上げます。\n"
        "　この度、一身上の都合により、退職届を同封いたしましたので、"
        "ご査収のほどよろしくお願い申し上げます。\n"
        "　在職中は格別のご厚情を賜り、誠にありがとうございました。\n敬具"
    )
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("記\n\n　・退職届　1通\n\n以上")
    run.font.name = "游明朝"
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "taishokutodoke_a4_yokogaki.docx")


def gen_taishokutodoke_b5_tategaki():
    """B5縦書き退職届"""
    doc = Document()
    set_a4_portrait(doc.sections[0], b5=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("退　職　届")
    run.font.name = "ＭＳ 明朝"
    run.font.size = Pt(24)
    run.bold = True

    body = (
        "",
        "　私事、",
        "",
        "　この度、一身上の都合により、",
        "　令和　年　月　日をもって退職いたします。",
        "",
        "　令和　年　月　日",
        "",
        "　　　　所属：　　　　　　　　　",
        "　　　　氏名：　　　　　　　　　㊞",
        "",
        "",
        "○○株式会社",
        "代表取締役　○○　○○　殿",
    )
    for line in body:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "ＭＳ 明朝"
        run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "taishokutodoke_b5_tategaki.docx")


def gen_taishokutodoke_xlsx():
    """退職届 Excel版（差し込み印刷用 + 一覧管理）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "退職届"

    thick = Border(
        left=Side(border_style='thick', color='000000'),
        right=Side(border_style='thick', color='000000'),
        top=Side(border_style='thick', color='000000'),
        bottom=Side(border_style='thick', color='000000'),
    )

    ws['B2'] = "退　職　届"
    ws['B2'].font = Font(name='游明朝', size=24, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center', vertical='center')
    ws.merge_cells('B2:F2')
    ws.row_dimensions[2].height = 40

    ws['B5'] = "提出日："
    ws['C5'] = "令和　年　月　日"

    ws['B7'] = "宛先："
    ws['C7'] = "○○株式会社　代表取締役　○○　○○　殿"

    ws['B10'] = "所属："
    ws['C10'] = ""
    ws['B11'] = "氏名："
    ws['C11'] = ""

    ws['B14'] = "退職理由："
    ws['C14'] = "一身上の都合"

    ws['B16'] = "退職日："
    ws['C16'] = "令和　年　月　日"

    ws['B19'] = "本文："
    ws['B20'] = "私事、この度、一身上の都合により、令和　年　月　日をもって退職いたします。"
    ws.merge_cells('B20:F20')

    # 列幅
    ws.column_dimensions['B'].width = 14
    for col in ['C', 'D', 'E', 'F']:
        ws.column_dimensions[col].width = 14

    # ノート
    note_ws = wb.create_sheet("使い方")
    note_ws['A1'] = "退職届テンプレート（Excel版）使い方"
    note_ws['A1'].font = Font(size=14, bold=True)
    note_ws['A3'] = "1. C5・C7・C10・C11・C16 のセルに必要事項を記入してください"
    note_ws['A4'] = "2. 必要に応じて B20 の本文を編集してください"
    note_ws['A5'] = "3. 印刷プレビューで A4 縦に収まるか確認してください"

    wb.save(OUT_DIR / "taishokutodoke.xlsx")


# ==========================================================
# T011: 労働条件通知書
# ==========================================================
def gen_roudoujouken_seishain():
    """労働条件通知書 正社員用"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("労働条件通知書")
    run.font.name = "游ゴシック"
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("殿")
    run.font.size = Pt(11)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("事業所名称・所在地：\n使用者職氏名：")
    run.font.size = Pt(11)

    doc.add_paragraph("")

    table = doc.add_table(rows=15, cols=2)
    table.style = "Table Grid"
    items = [
        ("契約期間", "期間の定めなし（無期労働契約）"),
        ("就業場所", "（雇入れ直後）　　　／（変更の範囲）　　　"),
        ("従事すべき業務", "（雇入れ直後）　　　／（変更の範囲）　　　"),
        ("始業・終業の時刻", "始業 〇時〇分／終業 〇時〇分"),
        ("休憩時間", "〇時〇分から〇時〇分までの〇分"),
        ("所定外労働", "有（時間外労働あり）／無"),
        ("休日", "毎週〇曜日、国民の祝日、その他（　　　　）"),
        ("休暇", "年次有給休暇： 法定通り付与／その他の休暇"),
        ("賃金", "基本給： 月給〇〇〇,〇〇〇円\n諸手当： （内訳）"),
        ("固定残業代の有無", "有（〇円・〇時間相当）／無"),
        ("賃金締切日・支払日", "毎月〇日締切・翌月〇日支払"),
        ("昇給", "有（時期：　／基準：　）／無"),
        ("退職に関する事項", "定年制：有（〇歳）／継続雇用制度：有\n自己都合退職： 退職する〇日前までに届出\n解雇事由： 就業規則による"),
        ("社会保険等の加入状況", "健康保険・厚生年金保険・雇用保険・労災保険"),
        ("その他", "詳細は就業規則による"),
    ]
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "游明朝"

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("※ 本通知書は労働基準法第15条に基づき交付するものです。\n※ 2024年4月改正により、就業場所・業務の「変更の範囲」明示が義務化されています。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_footer_note(doc)
    doc.save(OUT_DIR / "roudoujouken-tsuuchisho_seishain.docx")


def gen_roudoujouken_keiyaku():
    """労働条件通知書 契約社員用（有期）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("労働条件通知書（契約社員用）")
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph("")

    table = doc.add_table(rows=12, cols=2)
    table.style = "Table Grid"
    items = [
        ("契約期間", "期間の定めあり（〇年〇月〇日 〜 〇年〇月〇日）"),
        ("契約更新の有無", "自動更新／更新する場合がある／更新しない"),
        ("更新の上限", "通算契約期間〇年・更新回数〇回"),
        ("無期転換ルール", "本契約期間中に通算5年を超える場合、無期転換申込権が発生"),
        ("就業場所", "（雇入れ直後）　／（変更の範囲）　"),
        ("従事すべき業務", "（雇入れ直後）　／（変更の範囲）　"),
        ("始業・終業時刻", "始業 〇時〇分／終業 〇時〇分"),
        ("休日・休暇", "毎週〇曜日／年次有給休暇 法定通り"),
        ("賃金", "時給〇,〇〇〇円／月給〇〇〇,〇〇〇円"),
        ("賃金支払", "毎月〇日締切・翌月〇日支払"),
        ("退職", "契約期間満了で終了／自己都合退職可"),
        ("社会保険", "健康保険・厚生年金・雇用保険・労災保険"),
    ]
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "roudoujouken-tsuuchisho_keiyaku.docx")


def gen_roudoujouken_arbeit():
    """労働条件通知書 アルバイト/パート用"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("労働条件通知書（アルバイト・パート用）")
    run.font.size = Pt(20)
    run.bold = True

    doc.add_paragraph("")

    table = doc.add_table(rows=10, cols=2)
    table.style = "Table Grid"
    items = [
        ("契約期間", "期間の定めあり（〇年〇月〇日 〜 〇年〇月〇日）"),
        ("契約更新", "あり（更新条件：勤務態度・業務成績）／なし"),
        ("就業場所", "（雇入れ直後）　／（変更の範囲）　"),
        ("業務内容", "（例）販売業務／レジ業務／清掃／調理補助"),
        ("勤務時間", "シフト制（週〇日・1日〇時間程度）"),
        ("時給", "〇,〇〇〇円／時"),
        ("賃金支払", "毎月〇日締切・翌月〇日支払"),
        ("交通費", "有（実費・上限月〇,〇〇〇円）／無"),
        ("有給休暇", "雇入れの日から起算して6ヶ月継続勤務・所定労働日数の8割以上出勤で付与"),
        ("社会保険", "週20時間以上の場合：雇用保険／週30時間以上：健康保険・厚生年金"),
    ]
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "roudoujouken-tsuuchisho_arbeit.docx")


def gen_roudoujouken_xlsx():
    """労働条件通知書 Excel版"""
    wb = Workbook()
    ws = wb.active
    ws.title = "労働条件通知書"

    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill('solid', fgColor='F0F0F0')

    ws['B2'] = "労働条件通知書"
    ws['B2'].font = Font(size=18, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('B2:E2')

    ws['D4'] = "通知日："
    ws['E4'] = "令和　年　月　日"

    items = [
        ("契約期間", "（無期 / 有期：〇年〇月〇日〜〇年〇月〇日）"),
        ("就業場所（雇入れ直後）", ""),
        ("就業場所（変更の範囲）", ""),
        ("業務内容（雇入れ直後）", ""),
        ("業務内容（変更の範囲）", ""),
        ("始業時刻", ""),
        ("終業時刻", ""),
        ("休憩時間（分）", ""),
        ("所定休日", "毎週〇曜日・国民の祝日"),
        ("年次有給休暇", "法定通り付与"),
        ("基本給（月給）", ""),
        ("諸手当", ""),
        ("固定残業代", "有（〇円・〇時間相当）／無"),
        ("賃金締切日", "毎月〇日"),
        ("賃金支払日", "翌月〇日"),
        ("昇給", "有 / 無"),
        ("退職金", "有（規程あり）／無"),
        ("解雇事由", "就業規則による"),
        ("社会保険", "健康保険・厚生年金・雇用保険・労災保険"),
        ("無期転換申込権", "5年超で発生（有期契約者のみ）"),
    ]
    start_row = 6
    ws.cell(start_row, 2, "項目").font = Font(bold=True)
    ws.cell(start_row, 2).fill = header_fill
    ws.cell(start_row, 2).border = border
    ws.cell(start_row, 3, "内容").font = Font(bold=True)
    ws.cell(start_row, 3).fill = header_fill
    ws.cell(start_row, 3).border = border
    ws.merge_cells(start_row=start_row, start_column=3, end_row=start_row, end_column=5)

    for i, (k, v) in enumerate(items):
        r = start_row + 1 + i
        ws.cell(r, 2, k).border = border
        ws.cell(r, 3, v).border = border
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=5)

    ws.column_dimensions['B'].width = 28
    for col in ['C', 'D', 'E']:
        ws.column_dimensions[col].width = 16

    note = wb.create_sheet("使い方")
    note['A1'] = "労働条件通知書（Excel版）使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. C列の各項目に必要事項を入力してください"
    note['A4'] = "2. 2024年4月改正により『就業場所・業務の変更の範囲』を必ず明示してください"
    note['A5'] = "3. 印刷時は B〜E 列が用紙幅に収まるよう設定してください"
    note['A7'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A7'].font = Font(color='808080')

    wb.save(OUT_DIR / "roudoujouken-tsuuchisho.xlsx")


# ==========================================================
# T016: 業務委託契約書
# ==========================================================
def _gyoumu_itaku_common(doc, contract_type_label):
    """業務委託契約書 共通本文（タイプにより微調整）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"業務委託契約書（{contract_type_label}）")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run(
        "○○株式会社（以下「甲」という）と○○○○（以下「乙」という）は、"
        "甲が乙に対し業務を委託することにつき、以下のとおり契約（以下「本契約」という）を締結する。"
    )
    run.font.size = Pt(10)

    clauses = [
        ("第1条（業務内容）", "甲は乙に対し、以下の業務（以下「本業務」という）を委託し、乙はこれを受託する。\n（1）　　　　　　　　　　\n（2）　　　　　　　　　　"),
        ("第2条（契約期間）", "本契約の有効期間は、令和　年　月　日から令和　年　月　日までとする。期間満了の30日前までにいずれの当事者からも書面による解約の申出がない場合、本契約は同一条件でさらに1年間更新されるものとし、以後も同様とする。"),
        ("第3条（報酬）", "甲は乙に対し、本業務の対価として月額金〇〇〇,〇〇〇円（消費税別）を支払う。乙は毎月末日締めで甲に対して請求書を発行し、甲は翌月末日までに乙の指定する銀行口座に振り込む方法により支払う（フリーランス保護新法に基づき、業務完了から60日以内の支払いを保証）。"),
        ("第4条（業務遂行）", "乙は善良な管理者の注意をもって本業務を遂行する。乙は本業務の遂行方法、時間、場所について自己の裁量で決定するものとし、甲は乙に対して具体的な指揮命令を行わない。"),
        ("第5条（知的財産権）", "本業務の遂行に伴い創出された成果物の知的財産権は、甲が乙に対し本契約に基づく対価を全額支払った時点で乙から甲に移転する。ただし、乙が本契約締結前から有する知的財産権はこの限りでない。"),
        ("第6条（秘密保持）", "甲および乙は、本契約に関連して知り得た相手方の秘密情報を、相手方の事前の書面による承諾なく第三者に開示・漏洩してはならない。本条の義務は本契約終了後3年間有効とする。"),
        ("第7条（損害賠償）", "甲または乙が本契約上の義務に違反し、相手方に損害を与えた場合、その損害（直接損害に限る）を賠償する責任を負う。賠償額の上限は、本契約に基づき甲が乙に支払う直近12ヶ月分の報酬総額を超えないものとする。"),
        ("第8条（解除）", "甲または乙は、相手方が本契約上の義務に違反し、相当期間を定めた書面による催告をしても是正されない場合、本契約を解除することができる。"),
        ("第9条（インボイス）", "乙が適格請求書発行事業者である場合、登録番号を請求書に記載するものとする。乙が免税事業者の場合、消費税の取扱いについて別途協議する。"),
        ("第10条（反社会的勢力の排除）", "甲および乙は、自らおよびその役員・従業員が反社会的勢力に該当しないこと、反社会的勢力との関係を有しないことを表明・保証する。"),
        ("第11条（協議事項）", "本契約に定めのない事項または本契約の解釈に疑義が生じた場合は、甲乙誠意をもって協議のうえ解決する。"),
        ("第12条（準拠法・管轄）", "本契約は日本法に準拠し、本契約に関する一切の紛争については、〇〇地方裁判所を第一審の専属的合意管轄裁判所とする。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)

        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ各1通を保有する。\n\n"
        "令和　年　月　日"
    )
    run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("（甲）　　　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("（乙）　　　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)


def gen_gyoumu_itaku_general():
    """業務委託契約書 一般委託版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    _gyoumu_itaku_common(doc, "一般委託")
    add_footer_note(doc)
    doc.save(OUT_DIR / "gyoumu-itaku_general.docx")


def gen_gyoumu_itaku_junin():
    """業務委託契約書 準委任版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    _gyoumu_itaku_common(doc, "準委任")
    add_footer_note(doc)
    doc.save(OUT_DIR / "gyoumu-itaku_junin.docx")


def gen_gyoumu_itaku_ukeoi():
    """業務委託契約書 請負版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    _gyoumu_itaku_common(doc, "請負")
    add_footer_note(doc)
    doc.save(OUT_DIR / "gyoumu-itaku_ukeoi.docx")


# ==========================================================
# T002: 退職届 無料特化（T001 派生）
# ==========================================================

def gen_taishokutodoke_muryou_a4():
    """退職届 無料特化版・A4横書き"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("退　職　届")
    run.font.name = "游明朝"
    run.font.size = Pt(28)
    run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("所属：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "　私事、この度、一身上の都合により、令和　年　月　日をもって"
        "退職いたします。\n\n　以上"
    )
    run.font.size = Pt(11)
    add_footer_note(doc)
    doc.save(OUT_DIR / "taishokutodoke-muryou_a4.docx")


def gen_taishokutodoke_muryou_a4_tategaki():
    """退職届 無料特化版・A4縦書き"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    section = doc.sections[0]
    sectPr = section._sectPr
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'tbRl')
    sectPr.append(textDirection)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("退　職　届")
    run.font.name = "ＭＳ 明朝"
    run.font.size = Pt(28)
    run.bold = True
    body = (
        "", "", "　私事、", "",
        "　この度、一身上の都合により、",
        "　令和　年　月　日をもって退職いたします。",
        "", "", "　令和　年　月　日", "",
        "　　　　　　　　　所属：　　　　　　　　",
        "　　　　　　　　　氏名：　　　　　　　　　㊞",
        "", "", "○○株式会社", "代表取締役　○○　○○　殿",
    )
    for line in body:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "ＭＳ 明朝"
        run.font.size = Pt(11)
    add_footer_note(doc)
    doc.save(OUT_DIR / "taishokutodoke-muryou_a4_tategaki.docx")


# ==========================================================
# T007: 始末書テンプレート
# ==========================================================

def _shimatsusho_common(doc, title_kind, body_text):
    """始末書 共通フォーマット"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("始　末　書")
    run.font.name = "游明朝"
    run.font.size = Pt(24)
    run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("所属：　　　　　　　　　\n氏名：　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(body_text)
    run.font.size = Pt(11)


def gen_shimatsusho_general():
    """始末書 汎用版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    body = (
        "　この度、私の不注意により下記のような事態を引き起こしましたこと、深くお詫び申し上げます。\n\n"
        "【事案概要】\n　令和　年　月　日　　　　　　　　\n　　　　　　　　　　　　　　　　　\n\n"
        "【発生原因】\n　　　　　　　　　　　　　　　　　\n\n"
        "【再発防止策】\n　　　　　　　　　　　　　　　　　\n\n"
        "今後はこのようなことが二度と起こらぬよう、十分に注意して業務に従事することを誓います。\n\n以上"
    )
    _shimatsusho_common(doc, "汎用", body)
    add_footer_note(doc)
    doc.save(OUT_DIR / "shimatsusho_general.docx")


def gen_shimatsusho_chikoku():
    """始末書 遅刻特化"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    body = (
        "　この度、私の不注意により下記のとおり遅刻いたしましたこと、深くお詫び申し上げます。\n\n"
        "【遅刻日時】令和　年　月　日　　時　分到着\n\n"
        "【遅刻理由】\n　　　　　　　　　　　　　　　　　\n\n"
        "【再発防止策】\n　今後は始業時刻の30分前までに出勤するよう、生活習慣を改善いたします。\n\n"
        "業務に支障をきたしましたこと重ねてお詫び申し上げますとともに、二度と同様の事態を起こさぬよう、注意して業務に従事することを誓います。\n\n以上"
    )
    _shimatsusho_common(doc, "遅刻", body)
    add_footer_note(doc)
    doc.save(OUT_DIR / "shimatsusho_chikoku.docx")


def gen_shimatsusho_mudankin():
    """始末書 無断欠勤特化"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    body = (
        "　この度、私の不徳の致すところにより、下記の通り無断にて欠勤いたしましたこと、深くお詫び申し上げます。\n\n"
        "【欠勤日】令和　年　月　日（　曜日）\n\n"
        "【欠勤理由】\n　　　　　　　　　　　　　　　　　\n\n"
        "【再発防止策】\n　体調管理を徹底し、緊急時には始業時刻前に必ず連絡を入れるよう徹底いたします。\n\n"
        "勝手な行動により多大なご迷惑をおかけしましたこと、重ねてお詫び申し上げます。今後二度とこのようなことが起こらぬよう、誠心誠意業務に従事することを誓います。\n\n以上"
    )
    _shimatsusho_common(doc, "無断欠勤", body)
    add_footer_note(doc)
    doc.save(OUT_DIR / "shimatsusho_mudankin.docx")


def gen_shimatsusho_xlsx():
    """始末書 Excel版"""
    wb = Workbook()
    ws = wb.active
    ws.title = "始末書"
    ws['B2'] = "始　末　書"
    ws['B2'].font = Font(name='游明朝', size=22, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('B2:F2')
    ws['B5'] = "提出日："
    ws['C5'] = "令和　年　月　日"
    ws['B7'] = "宛先："
    ws['C7'] = "○○株式会社　代表取締役　○○　○○　殿"
    ws['B10'] = "所属："
    ws['B11'] = "氏名："
    ws['B14'] = "事案概要："
    ws['B15'] = "（具体的に記入）"
    ws.merge_cells('B15:F15')
    ws['B17'] = "発生原因："
    ws.merge_cells('B17:F17')
    ws['B19'] = "再発防止策："
    ws.merge_cells('B19:F19')
    ws.column_dimensions['B'].width = 14
    for col in ['C','D','E','F']: ws.column_dimensions[col].width = 14
    note = wb.create_sheet("使い方")
    note['A1'] = "始末書テンプレート（Excel版）使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. C5・C7・B11 のセルに必要事項を記入してください"
    note['A4'] = "2. 「事案概要」「発生原因」「再発防止策」を具体的に記入"
    note['A5'] = "3. 印刷プレビューで A4 縦に収まるか確認"
    note['A7'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A7'].font = Font(color='808080')
    wb.save(OUT_DIR / "shimatsusho.xlsx")


# ==========================================================
# T012: 雇用契約書テンプレート
# ==========================================================

def _koyou_keiyakusho_common(doc, kind_label, items):
    """雇用契約書 共通"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"雇用契約書（{kind_label}）")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "○○株式会社（以下「甲」という）と○○○○（以下「乙」という）は、"
        "甲が乙を雇用することにつき、以下のとおり雇用契約を締結する。"
    )
    run.font.size = Pt(10)
    doc.add_paragraph("")
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("本契約の成立を証するため、本書2通を作成し、甲乙各自署名押印のうえ、各1通を保有する。\n\n令和　年　月　日")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("（甲）　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("（乙）　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)


def gen_koyou_keiyakusho_seishain():
    """雇用契約書 正社員用"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("契約期間", "期間の定めなし（無期労働契約）"),
        ("試用期間", "採用日から　ヶ月（賃金条件は本契約と同じ）"),
        ("就業場所", "（雇入れ直後）／（変更の範囲）"),
        ("従事業務", "（雇入れ直後）／（変更の範囲）"),
        ("始業終業", "始業 　時　分／終業 　時　分／休憩 　分"),
        ("休日休暇", "毎週　曜日、国民の祝日／年次有給休暇 法定通り付与"),
        ("賃金", "基本給 月給　　　円／諸手当（内訳：　　　）"),
        ("固定残業代", "有（　円・　時間相当）／無"),
        ("賃金支払", "毎月　日締切・翌月　日支払（口座振込）"),
        ("退職", "定年制 　歳／自己都合：退職する　日前に届出"),
        ("社会保険", "健康保険・厚生年金・雇用保険・労災保険"),
        ("競業避止義務", "在職中および退職後　年間、競業他社への就業を禁止"),
        ("秘密保持", "本契約に関連して知り得た秘密情報を第三者に開示しない"),
    ]
    _koyou_keiyakusho_common(doc, "正社員", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "koyou-keiyakusho_seishain.docx")


def gen_koyou_keiyakusho_keiyaku():
    """雇用契約書 契約社員用"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("契約期間", "期間の定めあり（　年　月　日 〜 　年　月　日）"),
        ("契約更新", "自動更新／更新する場合がある／更新しない"),
        ("更新の上限", "通算契約期間 　年・更新回数 　回"),
        ("無期転換ルール", "5年超で無期転換申込権が発生（労契法18条）"),
        ("就業場所", "（雇入れ直後）／（変更の範囲）"),
        ("従事業務", "（雇入れ直後）／（変更の範囲）"),
        ("始業終業", "始業 　時　分／終業 　時　分"),
        ("休日休暇", "毎週　曜日／年次有給休暇 法定通り"),
        ("賃金", "時給／月給 　　　円"),
        ("賃金支払", "毎月　日締切・翌月　日支払"),
        ("退職", "契約期間満了で終了"),
        ("社会保険", "健康保険・厚生年金・雇用保険・労災保険"),
    ]
    _koyou_keiyakusho_common(doc, "契約社員", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "koyou-keiyakusho_keiyaku.docx")


def gen_koyou_keiyakusho_arbeit():
    """雇用契約書 アルバイト/パート用"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("契約期間", "期間の定めあり（　年　月　日 〜 　年　月　日）"),
        ("契約更新", "あり（更新条件：勤務態度・業務成績）／なし"),
        ("就業場所", "（雇入れ直後）／（変更の範囲）"),
        ("業務内容", "（例）販売／レジ／清掃／調理補助"),
        ("勤務時間", "シフト制（週　日・1日　時間程度）"),
        ("時給", "　　　円／時"),
        ("賃金支払", "毎月　日締切・翌月　日支払"),
        ("交通費", "有（実費・上限月　円）／無"),
        ("有給休暇", "雇入れ6ヶ月後・所定労働日数の8割以上出勤で付与"),
        ("社会保険", "週20時間以上：雇用保険／週30時間以上：健康保険・厚生年金"),
    ]
    _koyou_keiyakusho_common(doc, "アルバイト・パート", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "koyou-keiyakusho_arbeit.docx")


def gen_koyou_keiyakusho_shokutaku():
    """雇用契約書 嘱託用（再雇用）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("契約期間", "期間の定めあり（　年　月　日 〜 　年　月　日）"),
        ("身分", "嘱託社員（定年後再雇用）"),
        ("就業場所", "（雇入れ直後）／（変更の範囲）"),
        ("業務内容", "従前の業務と同等または軽減した業務"),
        ("勤務時間", "始業 　時　分／終業 　時　分（短縮可）"),
        ("休日", "毎週　曜日／年次有給休暇 法定通り"),
        ("賃金", "月給 　　　円（在職時の　%目安）"),
        ("賞与", "有 / 無"),
        ("社会保険", "健康保険・厚生年金（70歳まで）・雇用保険"),
        ("退職", "契約期間満了 / 65歳まで再雇用可"),
    ]
    _koyou_keiyakusho_common(doc, "嘱託", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "koyou-keiyakusho_shokutaku.docx")


# ==========================================================
# T023: 請求書（個人事業主）
# ==========================================================
def gen_seikyu_kojin_xlsx():
    """請求書 個人事業主向け Excel版（インボイス対応・自動計算）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "請求書"
    thin = Side(border_style='thin', color='000000')
    medium = Side(border_style='medium', color='000000')

    ws['F2'] = "請　求　書"
    ws['F2'].font = Font(name='游明朝', size=22, bold=True)
    ws['F2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('F2:I2')

    ws['B5'] = "請求先："
    ws['B5'].font = Font(bold=True)
    ws['C5'] = ""
    ws['C5'].font = Font(size=12)
    ws['B6'] = ""

    ws['F5'] = "請求日："
    ws['F5'].font = Font(bold=True)
    ws['G5'] = "令和　年　月　日"
    ws['F6'] = "請求番号："
    ws['F6'].font = Font(bold=True)
    ws['G6'] = ""
    ws['F7'] = "支払期限："
    ws['F7'].font = Font(bold=True)
    ws['G7'] = "令和　年　月　日"

    ws['B9'] = "下記の通りご請求申し上げます。"
    ws['B9'].font = Font(size=11)

    # 明細表
    headers = ["項目", "数量", "単価", "金額"]
    for i, h in enumerate(headers):
        c = ws.cell(11, 2+i, h)
        c.font = Font(bold=True)
        c.fill = PatternFill('solid', fgColor='F0F0F0')
        c.border = Border(left=thin, right=thin, top=medium, bottom=medium)
        c.alignment = Alignment(horizontal='center')
    # 明細行
    for r in range(12, 22):
        for col in range(2, 6):
            ws.cell(r, col, "" if col != 5 else None).border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # 自動計算式
    for r in range(12, 22):
        ws.cell(r, 5).value = f'=IF(AND(C{r}<>"",D{r}<>""),C{r}*D{r},"")'

    # 集計
    ws['D23'] = "小計"; ws['D23'].font = Font(bold=True); ws['D23'].alignment = Alignment(horizontal='right')
    ws['E23'] = "=SUM(E12:E21)"; ws['E23'].number_format = '#,##0'
    ws['D24'] = "消費税(10%)"; ws['D24'].font = Font(bold=True); ws['D24'].alignment = Alignment(horizontal='right')
    ws['E24'] = "=ROUNDDOWN(E23*0.1,0)"; ws['E24'].number_format = '#,##0'
    ws['D25'] = "源泉徴収税(10.21%)"; ws['D25'].font = Font(bold=True); ws['D25'].alignment = Alignment(horizontal='right')
    ws['E25'] = "=ROUNDDOWN(E23*0.1021,0)"; ws['E25'].number_format = '#,##0'
    ws['D26'] = "合計"; ws['D26'].font = Font(bold=True, size=12); ws['D26'].alignment = Alignment(horizontal='right')
    ws['E26'] = "=E23+E24-E25"; ws['E26'].number_format = '#,##0'; ws['E26'].font = Font(bold=True, size=12)

    # 振込先
    ws['B28'] = "振込先"
    ws['B28'].font = Font(bold=True)
    ws['B29'] = "銀行名："
    ws['B30'] = "支店名："
    ws['B31'] = "口座種別："
    ws['B32'] = "口座番号："
    ws['B33'] = "口座名義："
    ws['B34'] = "適格請求書発行事業者番号（インボイス）："

    # 発行者
    ws['F28'] = "発行者"
    ws['F28'].font = Font(bold=True)
    ws['F29'] = "氏名／屋号："
    ws['F30'] = "住所："
    ws['F31'] = "電話："
    ws['F32'] = "メール："

    # 列幅
    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 8
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 18
    for col in ['G','H','I']: ws.column_dimensions[col].width = 14

    note = wb.create_sheet("使い方")
    note['A1'] = "請求書テンプレート（個人事業主向け・Excel版）使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. 請求先・請求日・請求番号・支払期限を記入"
    note['A4'] = "2. 明細欄（B12〜D21）に項目・数量・単価を記入 → 金額は自動計算"
    note['A5'] = "3. 小計・消費税・源泉徴収税・合計は数式で自動計算（B23-E26）"
    note['A6'] = "4. 振込先（B29-B33）と適格請求書発行事業者番号（B34）を記入"
    note['A7'] = "5. 印刷時は A4 縦に収まるよう設定"
    note['A9'] = "源泉徴収税について："
    note['A10'] = "  デザイン・ライティング業務など特定の業務は10.21%源泉徴収"
    note['A11'] = "  該当しない場合は E25 を 0 にする"
    note['A13'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A13'].font = Font(color='808080')

    wb.save(OUT_DIR / "seikyu-kojin-jigyonushi.xlsx")


def gen_seikyu_kojin_word():
    """請求書 個人事業主向け Word版（参考用）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("請　求　書")
    run.font.name = "游明朝"
    run.font.size = Pt(24)
    run.bold = True
    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "請求先"
    table.cell(0, 1).text = "　　　　　　　御中"
    table.cell(1, 0).text = "請求日"
    table.cell(1, 1).text = "令和　年　月　日"
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("下記の通りご請求申し上げます。")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    table2 = doc.add_table(rows=11, cols=4)
    table2.style = "Table Grid"
    headers = ["項目", "数量", "単価", "金額"]
    for i, h in enumerate(headers):
        cell = table2.cell(0, i)
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for r in range(1, 11):
        for col in range(4):
            table2.cell(r, col).text = ""

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("小計：　　　　　　円\n消費税（10%）：　　　　　　円\n源泉徴収税（10.21%）：　△　　　　　　円\n合計：　　　　　　円")
    run.font.size = Pt(11)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("【振込先】\n銀行名：\n支店名：\n口座種別：\n口座番号：\n口座名義：\n適格請求書発行事業者番号（インボイス）：T\n\n【発行者】\n氏名／屋号：\n住所：\n電話：")
    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "seikyu-kojin-jigyonushi.docx")


# ==========================================================
# T029: 見積書（登録不要）
# ==========================================================
def gen_mitsumori_xlsx():
    """見積書 Excel版（自動計算）"""
    wb = Workbook()
    ws = wb.active
    ws.title = "見積書"
    thin = Side(border_style='thin', color='000000')
    medium = Side(border_style='medium', color='000000')

    ws['F2'] = "見　積　書"
    ws['F2'].font = Font(name='游明朝', size=22, bold=True)
    ws['F2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('F2:I2')

    ws['B5'] = "宛先："; ws['B5'].font = Font(bold=True)
    ws['F5'] = "見積日："; ws['F5'].font = Font(bold=True); ws['G5'] = "令和　年　月　日"
    ws['F6'] = "見積番号："; ws['F6'].font = Font(bold=True)
    ws['F7'] = "有効期限："; ws['F7'].font = Font(bold=True); ws['G7'] = "発行日より　日間"

    ws['B9'] = "下記の通りお見積申し上げます。"

    headers = ["項目", "数量", "単価", "金額"]
    for i, h in enumerate(headers):
        c = ws.cell(11, 2+i, h)
        c.font = Font(bold=True)
        c.fill = PatternFill('solid', fgColor='F0F0F0')
        c.border = Border(left=thin, right=thin, top=medium, bottom=medium)
        c.alignment = Alignment(horizontal='center')
    for r in range(12, 22):
        for col in range(2, 6):
            ws.cell(r, col).border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for r in range(12, 22):
        ws.cell(r, 5).value = f'=IF(AND(C{r}<>"",D{r}<>""),C{r}*D{r},"")'

    ws['D23'] = "小計"; ws['D23'].font = Font(bold=True); ws['D23'].alignment = Alignment(horizontal='right')
    ws['E23'] = "=SUM(E12:E21)"; ws['E23'].number_format = '#,##0'
    ws['D24'] = "消費税(10%)"; ws['D24'].font = Font(bold=True); ws['D24'].alignment = Alignment(horizontal='right')
    ws['E24'] = "=ROUNDDOWN(E23*0.1,0)"; ws['E24'].number_format = '#,##0'
    ws['D25'] = "合計"; ws['D25'].font = Font(bold=True, size=12); ws['D25'].alignment = Alignment(horizontal='right')
    ws['E25'] = "=E23+E24"; ws['E25'].number_format = '#,##0'; ws['E25'].font = Font(bold=True, size=12)

    ws['B27'] = "発行者"
    ws['B27'].font = Font(bold=True)
    ws['B28'] = "会社名／屋号："
    ws['B29'] = "担当者："
    ws['B30'] = "住所："
    ws['B31'] = "電話："
    ws['B32'] = "メール："
    ws['B33'] = "備考："

    ws.column_dimensions['B'].width = 28
    ws.column_dimensions['C'].width = 8
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 18
    for col in ['G','H','I']: ws.column_dimensions[col].width = 14

    note = wb.create_sheet("使い方")
    note['A1'] = "見積書テンプレート Excel版 使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. 宛先・見積日・有効期限を記入"
    note['A4'] = "2. 明細欄に項目・数量・単価を記入 → 金額は自動計算"
    note['A5'] = "3. 発行者情報を記入"
    note['A6'] = "4. 印刷時は A4 縦に収まるよう設定"
    note['A8'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A8'].font = Font(color='808080')

    wb.save(OUT_DIR / "mitsumori-tourokufuyou.xlsx")


def gen_mitsumori_word():
    """見積書 Word版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("見　積　書")
    run.font.name = "游明朝"
    run.font.size = Pt(24)
    run.bold = True
    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "宛先"
    table.cell(0, 1).text = "　　　　　　　御中"
    table.cell(1, 0).text = "見積日"
    table.cell(1, 1).text = "令和　年　月　日"

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("下記の通りお見積申し上げます。")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    table2 = doc.add_table(rows=11, cols=4)
    table2.style = "Table Grid"
    headers = ["項目", "数量", "単価", "金額"]
    for i, h in enumerate(headers):
        cell = table2.cell(0, i)
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for r in range(1, 11):
        for col in range(4):
            table2.cell(r, col).text = ""

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("小計：　　　　　　円\n消費税（10%）：　　　　　　円\n合計：　　　　　　円")
    run.font.size = Pt(11)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("【発行者】\n会社名／屋号：\n担当者：\n住所：\n電話：\n備考：")
    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "mitsumori-tourokufuyou.docx")


# ==========================================================
# T035: 議事録
# ==========================================================
def _gijiroku_common(doc, title_label, items):
    """議事録 共通構造"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"議　事　録（{title_label}）")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")

    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v


def gen_gijiroku_standard():
    """議事録 標準版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("会議名", ""),
        ("日時", "令和　年　月　日（　）　時　分 〜 　時　分"),
        ("場所", ""),
        ("出席者", ""),
        ("欠席者", ""),
        ("議題", "1. \n2. \n3. "),
        ("議事内容", ""),
        ("決定事項", "・\n・"),
        ("ToDo / 担当者", "・　　　　　　（担当：　　／期限：　　）"),
        ("次回予定", ""),
        ("作成者", ""),
    ]
    _gijiroku_common(doc, "標準", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "gijiroku_standard.docx")


def gen_gijiroku_1on1():
    """議事録 1on1版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("実施日時", "令和　年　月　日（　）　時　分 〜 　時　分"),
        ("対象者", ""),
        ("実施者（マネージャー）", ""),
        ("前回からの進捗", ""),
        ("業務面の話題", ""),
        ("個人面（キャリア・成長）", ""),
        ("課題・困りごと", ""),
        ("マネージャーからのフィードバック", ""),
        ("ネクストアクション", "・"),
        ("次回予定", ""),
    ]
    _gijiroku_common(doc, "1on1", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "gijiroku_1on1.docx")


def gen_gijiroku_brest():
    """議事録 ブレスト版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    items = [
        ("テーマ", ""),
        ("日時", "令和　年　月　日（　）　時　分 〜 　時　分"),
        ("参加者", ""),
        ("ファシリテーター", ""),
        ("背景・目的", ""),
        ("出されたアイデア", "・\n・\n・\n・\n・"),
        ("評価・絞り込み", ""),
        ("採用案", ""),
        ("次のアクション", ""),
    ]
    _gijiroku_common(doc, "ブレスト", items)
    add_footer_note(doc)
    doc.save(OUT_DIR / "gijiroku_brest.docx")


def gen_gijiroku_xlsx():
    """議事録 Excel版"""
    wb = Workbook()
    ws = wb.active
    ws.title = "議事録"
    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill('solid', fgColor='F0F0F0')

    ws['B2'] = "議　事　録"
    ws['B2'].font = Font(size=18, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('B2:E2')

    items = [
        ("会議名", ""),
        ("日時", "令和　年　月　日"),
        ("場所", ""),
        ("出席者", ""),
        ("議題", ""),
        ("議事内容", ""),
        ("決定事項", ""),
        ("ToDo・担当者・期限", ""),
        ("次回予定", ""),
    ]
    start_row = 5
    for i, (k, v) in enumerate(items):
        r = start_row + i
        c1 = ws.cell(r, 2, k)
        c1.font = Font(bold=True)
        c1.fill = header_fill
        c1.border = border
        c1.alignment = Alignment(vertical='top')
        c2 = ws.cell(r, 3, v)
        c2.border = border
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=5)
        ws.row_dimensions[r].height = 35

    ws.column_dimensions['B'].width = 22
    for col in ['C','D','E']: ws.column_dimensions[col].width = 22

    note = wb.create_sheet("使い方")
    note['A1'] = "議事録テンプレート Excel版 使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. C列（マージ済）の各項目に内容を記入"
    note['A4'] = "2. 行の高さは内容に合わせて調整"
    note['A5'] = "3. 印刷時は A4 縦に収まるよう設定"
    note['A7'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A7'].font = Font(color='808080')

    wb.save(OUT_DIR / "gijiroku.xlsx")


# ==========================================================
# T020: 秘密保持誓約書（NDA）
# ==========================================================
def gen_himitsuhoji_nyusha():
    """秘密保持誓約書 雇用時版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("秘密保持誓約書（雇用時）")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    clauses = [
        ("第1条（秘密保持義務）", "私は、貴社における業務上知り得た下記の情報（以下「秘密情報」という）について、貴社の事前の書面による承諾なく、第三者に開示・漏洩しない。\n（1）顧客情報・取引先情報\n（2）財務情報・経営情報\n（3）人事情報・社員情報\n（4）技術情報・営業ノウハウ\n（5）その他、貴社が秘密として指定した情報"),
        ("第2条（目的外使用の禁止）", "私は、秘密情報を業務遂行以外の目的で使用しない。"),
        ("第3条（複製の禁止）", "秘密情報を含む書類・電子データは、業務上必要な範囲を超えて複製しない。"),
        ("第4条（情報の返還）", "退職時、または貴社の指示があったときは、秘密情報を含む一切の書類・電子データを直ちに返還または破棄する。"),
        ("第5条（損害賠償）", "本誓約に違反し、貴社に損害を与えた場合、その損害を賠償する。"),
        ("第6条（有効期間）", "本誓約による秘密保持義務は、退職後も継続する。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("以上の通り誓約いたします。")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "himitsuhoji-seiyakusho_nyusha.docx")


def gen_himitsuhoji_taisha():
    """秘密保持誓約書 退職時版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("秘密保持誓約書（退職時）")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    clauses = [
        ("第1条（情報の返還）", "私は、退職にあたり、在職中に取得した貴社の秘密情報・営業情報・顧客情報を含む一切の書類・電子データ・媒体を貴社に返還または破棄したことを確認する。"),
        ("第2条（秘密保持義務の継続）", "退職後も、在職中に知り得た秘密情報を第三者に開示・漏洩しない。"),
        ("第3条（競業避止義務）", "退職後　年間、貴社の同業他社に就業しない、または貴社と競合する事業を営まない。\n※ 期間・地域は合理的な範囲（最高裁判例に基づき　年程度）で個別協議"),
        ("第4条（顧客等への接触禁止）", "退職後　年間、在職中に取引した貴社の顧客・取引先に営業活動を行わない。"),
        ("第5条（損害賠償）", "本誓約に違反し、貴社に損害を与えた場合、その損害を賠償する。"),
        ("第6条（裁判管轄）", "本誓約に関する紛争は、貴社所在地を管轄する裁判所を専属的合意管轄裁判所とする。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("以上の通り誓約いたします。")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "himitsuhoji-seiyakusho_taisha.docx")


# ==========================================================
# T019: 秘密保持契約書（NDA）双方向 / 片務
# ==========================================================
def gen_nda_souhou():
    """秘密保持契約書（双方向NDA）— 取引先間の相互開示"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("秘密保持契約書")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社（以下「甲」という）と○○株式会社（以下「乙」という）は、両者間における取引（以下「本取引」という）に関連して開示する秘密情報の取扱いについて、次のとおり秘密保持契約（以下「本契約」という）を締結する。")
    run.font.size = Pt(10)
    doc.add_paragraph("")

    clauses = [
        ("第1条（秘密情報の定義）",
         "本契約において「秘密情報」とは、本取引に関連して甲または乙が相手方に開示する一切の情報のうち、書面または電子データにより秘密である旨を明示して開示されたもの、および口頭で開示された場合は開示後14日以内に書面で確認したものをいう。\nただし、次の各号に該当する情報は秘密情報に含まれない。\n（1）開示時に既に公知であった情報\n（2）開示後、受領者の責によらず公知となった情報\n（3）受領者が独自に開発した情報\n（4）正当な権限を有する第三者から守秘義務なく取得した情報"),
        ("第2条（秘密保持義務）",
         "甲および乙は、相互に開示された秘密情報について、相手方の事前の書面による承諾なく、第三者に開示または漏洩してはならない。"),
        ("第3条（目的外使用の禁止）",
         "甲および乙は、相手方から開示された秘密情報を本取引の遂行以外の目的で使用してはならない。"),
        ("第4条（複製の制限）",
         "甲および乙は、相手方から開示された秘密情報を、本取引の遂行に必要な範囲を超えて複製してはならない。"),
        ("第5条（情報の返還・廃棄）",
         "本契約終了時、または相手方から要求があったときは、開示された秘密情報を含む一切の書類・電子データ・媒体を、相手方の指示に従って返還または廃棄しなければならない。"),
        ("第6条（有効期間）",
         "本契約の有効期間は、本契約締結日から○年間とする。ただし、本契約終了後も第2条（秘密保持義務）および第3条（目的外使用の禁止）の規定は、本契約終了から○年間有効に存続する。"),
        ("第7条（損害賠償）",
         "甲または乙が本契約に違反し、相手方に損害を与えた場合、違反した当事者は相手方に対し、その損害を賠償する責を負う。"),
        ("第8条（合意管轄）",
         "本契約に関する紛争は、○○地方裁判所を第一審の専属的合意管轄裁判所とする。"),
        ("第9条（協議事項）",
         "本契約に定めなき事項、または本契約の解釈について疑義が生じた事項については、甲乙誠実に協議の上、解決するものとする。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("本契約締結の証として、本書2通を作成し、甲乙各1通を保有する。")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【甲】\n住所：　　　　　　　　　　　　　　\n会社名：　　　　　　　　　　　　　\n代表者：　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【乙】\n住所：　　　　　　　　　　　　　　\n会社名：　　　　　　　　　　　　　\n代表者：　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "nda_souhou.docx")


def gen_nda_henmu():
    """秘密保持契約書（片務NDA）— 一方が一方的に情報開示するケース"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("秘密保持契約書（片務）")
    run.font.size = Pt(20)
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社（以下「開示者」という）と○○株式会社（以下「受領者」という）は、開示者から受領者へ開示される秘密情報の取扱いについて、次のとおり秘密保持契約（以下「本契約」という）を締結する。")
    run.font.size = Pt(10)
    doc.add_paragraph("")

    clauses = [
        ("第1条（秘密情報の定義）",
         "本契約において「秘密情報」とは、開示者から受領者に開示される一切の情報のうち、書面・電子データ・口頭・有形媒体を問わず、開示者が秘密として指定した情報をいう。\nただし、次の各号に該当する情報は秘密情報に含まれない。\n（1）開示時に既に公知であった情報\n（2）開示後、受領者の責によらず公知となった情報\n（3）受領者が独自に開発した情報\n（4）正当な権限を有する第三者から守秘義務なく取得した情報"),
        ("第2条（秘密保持義務）",
         "受領者は、開示者から開示された秘密情報について、開示者の事前の書面による承諾なく、第三者に開示または漏洩してはならない。"),
        ("第3条（目的外使用の禁止）",
         "受領者は、開示者から開示された秘密情報を、本契約に定める目的以外に使用してはならない。"),
        ("第4条（複製の制限）",
         "受領者は、開示者から開示された秘密情報を、本契約に定める目的の達成に必要な範囲を超えて複製してはならない。"),
        ("第5条（情報の返還・廃棄）",
         "本契約終了時、または開示者から要求があったときは、受領者は開示された秘密情報を含む一切の書類・電子データ・媒体を、開示者の指示に従って返還または廃棄しなければならない。"),
        ("第6条（有効期間および存続）",
         "本契約の有効期間は、本契約締結日から○年間とする。ただし、本契約終了後も第2条（秘密保持義務）および第3条（目的外使用の禁止）の規定は、本契約終了から○年間有効に存続する。"),
        ("第7条（損害賠償）",
         "受領者が本契約に違反し、開示者に損害を与えた場合、その損害を賠償する責を負う。"),
        ("第8条（合意管轄）",
         "本契約に関する紛争は、開示者の所在地を管轄する地方裁判所を第一審の専属的合意管轄裁判所とする。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body)
        run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("本契約締結の証として、本書2通を作成し、開示者・受領者各1通を保有する。")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【開示者】\n住所：　　　　　　　　　　　　　　\n会社名：　　　　　　　　　　　　　\n代表者：　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【受領者】\n住所：　　　　　　　　　　　　　　\n会社名：　　　　　　　　　　　　　\n代表者：　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "nda_henmu.docx")


# ==========================================================
# T028: 見積依頼書（RFQ）
# ==========================================================
# ==========================================================
# Vol.2: T041 念書 / T054 日報 / T057 お礼状
# ==========================================================
def gen_nensho_general():
    """念書（汎用）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("念　書")
    run.font.size = Pt(22); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    body_lines = [
        ("第1条（債務の確認）", "私（以下「乙」という）は、甲（○○○○）に対し、下記の事項を約束いたします。"),
        ("第2条（履行の約束）", "・履行内容： （金額・期日・行為の内容を具体的に記載）\n・履行期限： 令和○年○月○日\n・履行方法： （振込・現金交付・実行行為等）"),
        ("第3条（違反時の措置）", "上記の履行を怠った場合、甲は本念書を証拠として法的措置（民事訴訟・債権回収手続）を執ることができることを了承します。"),
        ("第4条（公正証書化への協力）", "甲が必要と認めた場合は、本念書を公正証書として作成することに協力します。"),
        ("第5条（連帯保証）", "（任意・連帯保証人を立てる場合に記載） 連帯保証人 ○○○○は、本念書に定める義務について、乙と連帯して責任を負います。"),
    ]
    for title, body in body_lines:
        p = doc.add_paragraph()
        run = p.add_run(title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "nensho_general.docx")


def gen_nippo_standard():
    """日報（標準・業務日報）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("業 務 日 報")
    run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")

    # ヘッダー情報テーブル
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid'
    cells = [
        ("日付", "令和　年　月　日（　）", "氏名", "　　　　　　　　"),
        ("所属部署", "　　　　　　　　", "天候", "　　"),
        ("出勤時刻", "　：　", "退勤時刻", "　：　"),
    ]
    for i, row in enumerate(cells):
        for j, text in enumerate(row):
            table.rows[i].cells[j].text = text

    doc.add_paragraph("")

    sections = [
        "■ 本日の業務内容（5W1Hで簡潔に）",
        "1. \n2. \n3. \n",
        "■ 達成事項・成果",
        "・\n・\n",
        "■ 課題・問題点",
        "・\n・\n",
        "■ 明日以降の予定・対応事項",
        "・\n・\n",
        "■ 上司への確認事項・申送り",
        "・\n",
        "■ 所感・気付き",
        "",
    ]
    for line in sections:
        p = doc.add_paragraph()
        run = p.add_run(line)
        if line.startswith("■"):
            run.bold = True; run.font.size = Pt(11)
        else:
            run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "nippo_standard.docx")


def gen_nippo_xlsx():
    """日報 Excel（関数で自動集計）"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "日報"
    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 10
    ws.column_dimensions['E'].width = 14

    # タイトル
    ws['B2'] = "業務日報"
    ws['B2'].font = Font(size=18, bold=True)
    ws.merge_cells('B2:E2')

    # ヘッダー情報
    ws['B4'] = "日付"; ws['C4'] = "令和　年　月　日"
    ws['B5'] = "氏名"; ws['C5'] = ""
    ws['B6'] = "所属"; ws['C6'] = ""

    # タスクテーブル
    ws['B8'] = "No."; ws['C8'] = "業務内容"; ws['D8'] = "所要時間"; ws['E8'] = "成果・進捗"
    header_fill = PatternFill('solid', fgColor='E8F0FE')
    for col in ['B8', 'C8', 'D8', 'E8']:
        ws[col].font = Font(bold=True)
        ws[col].fill = header_fill
        ws[col].alignment = Alignment(horizontal='center')

    for i in range(9, 14):
        ws[f'B{i}'] = i - 8
        ws[f'B{i}'].alignment = Alignment(horizontal='center')

    # 合計行
    ws['B15'] = "合計"
    ws['B15'].font = Font(bold=True)
    ws['D15'] = "=SUM(D9:D14)"
    ws['D15'].font = Font(bold=True)

    # ノート
    ws['B17'] = "■ 課題・申送り事項"
    ws['B17'].font = Font(bold=True)
    ws.merge_cells('B18:E22')

    wb.save(OUT_DIR / "nippo.xlsx")


def gen_nensho_shiharai():
    """念書（支払い特化）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("念　書（支払い）")
    run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○　○○　殿"); run.font.size = Pt(11)
    doc.add_paragraph("")

    clauses = [
        ("第1条（債務の確認）", "私（以下「乙」という）は、貴殿（以下「甲」という）に対し、令和　年　月　日付けで発生した金 　万円の支払い債務を負っていることを確認します。"),
        ("第2条（支払いの約束）", "上記債務について、下記の方法で確実に支払うことを約束いたします。\n・支払い方法： 一括 または 分割 （該当に丸）\n・支払い期日： 令和　年　月　日\n・支払い場所： 甲指定の銀行口座への振込（振込手数料は乙負担）"),
        ("第3条（分割払いの場合）", "分割払いとする場合、下記のスケジュールに従って支払います。\n・第1回： 令和　年　月　日 　万円\n・第2回： 令和　年　月　日 　万円\n・第3回： 令和　年　月　日 　万円\n（必要に応じて行を追加）"),
        ("第4条（遅延損害金）", "支払期日に遅延した場合、年率14.6％の遅延損害金を甲に支払うことに同意します。"),
        ("第5条（期限の利益喪失）", "1回でも支払いを怠った場合、残債務全額について期限の利益を失い、直ちに一括して支払うこととします。"),
        ("第6条（公正証書化への協力）", "甲が必要と認めた場合は、本念書を公正証書として作成することに協力します。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    add_footer_note(doc)
    doc.save(OUT_DIR / "nensho_shiharai.docx")


def gen_nensho_kinsen():
    """念書（金銭・債務承認）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("念　書（金銭・債務承認）")
    run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○　○○　殿"); run.font.size = Pt(11)
    doc.add_paragraph("")

    clauses = [
        ("第1条（債務の承認）", "私（以下「乙」という）は、貴殿（以下「甲」という）に対し、下記金銭債務が存在することを承認いたします。\n・原契約日： 令和　年　月　日（金銭消費貸借）\n・原契約金額： 金 　万円\n・現在の残債務： 金 　万円\n・利息／遅延損害金： （該当に応じて記載）"),
        ("第2条（消滅時効の援用放棄）", "本債務について、民法上の消滅時効を援用しないことを誓約します。本書面による債務の承認をもって、民法第152条第1項に基づき時効の更新が生じることに同意します。"),
        ("第3条（返済の約束）", "上記残債務について、令和　年　月　日までに、甲指定の銀行口座へ全額一括にて返済します。\nまたは、別紙「返済計画書」に従い分割返済します。"),
        ("第4条（連帯保証）", "（任意）連帯保証人 ○○　○○ は、本念書記載の債務について乙と連帯して責任を負います。"),
        ("第5条（違反時の措置）", "返済期日までに支払いがない場合、甲は本念書を証拠として、訴訟・強制執行を含む法的措置を執ることに乙は異議なく同意します。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    add_footer_note(doc)
    doc.save(OUT_DIR / "nensho_kinsen.docx")


def gen_tenmatsusho():
    """顛末書（汎用）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("顛 末 書")
    run.font.size = Pt(22); run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n代表取締役　○○　○○　殿")
    run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("所属：　　　　　　　　　　　\n氏名：　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    sections = [
        ("件名", "（例：商品納品ミスに関する顛末書、業務上トラブル発生に関する顛末書 等）"),
        ("発生日時", "令和　年　月　日　午前／午後　時　分頃"),
        ("発生場所", "（営業先・社内・倉庫・現場 等を具体的に）"),
        ("関係者", "（自社：氏名・役職／取引先：会社名・氏名）"),
        ("経緯（時系列）", "1. 発生前の状況：\n2. 発生時の状況：\n3. 発生後の対応：\n4. 現時点の状況："),
        ("原因（直接原因・根本原因）", "・直接原因： （現象として何が起きたか）\n・根本原因： （なぜ発生したか・5なぜ分析）"),
        ("影響範囲", "・お客様への影響：\n・社内への影響：\n・金銭的損失（試算）："),
        ("再発防止策", "1. \n2. \n3. \n（業務プロセス改善・チェック体制強化・教育訓練 等を具体的に）"),
        ("結語", "今後はこのような事態を繰り返さぬよう、再発防止策を遵守し、業務にあたります。\nこの度は誠に申し訳ございませんでした。"),
    ]
    for title, body in sections:
        p = doc.add_paragraph()
        run = p.add_run("■ " + title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "tenmatsusho_general.docx")


def gen_nensho_kojin():
    """念書（個人間・夫婦/家族/友人）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("念　書（個人間）")
    run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○　○○　殿"); run.font.size = Pt(11)
    doc.add_paragraph("")

    clauses = [
        ("第1条（誓約事項）", "私（以下「乙」という）は、貴殿（以下「甲」という）に対し、下記の事項を約束いたします。\n（1）禁止行為： （例：浮気・金銭の無断引出し・暴言暴力 等を具体的に）\n（2）履行事項： （例：金銭の返済・養育費・慰謝料 等を具体的に）"),
        ("第2条（履行内容・期限）", "・履行内容： 金 　万円の支払い（または行為の停止）\n・履行期限： 令和　年　月　日\n・履行方法： 甲指定の銀行口座への振込（または直接交付）"),
        ("第3条（違反時の措置）", "上記の禁止事項を破り、または履行を怠った場合、下記の措置を受けることに同意します。\n（1）違約金 金 　万円を直ちに支払うこと\n（2）甲が本念書を証拠として、損害賠償請求・離婚調停・刑事告訴等の法的措置を執ること"),
        ("第4条（家族関係への配慮）", "本件が家族関係に与える影響を考慮し、両者は誠実に協議のうえ、家庭の平穏維持に努めることを確認します。"),
        ("第5条（公正証書化への協力）", "甲が必要と認めた場合は、本念書を公正証書として作成することに協力します。"),
        ("第6条（守秘）", "本念書の存在および内容を、両者は第三者に開示しないものとします（ただし、専門家への相談・法的手続きを除く）。"),
    ]
    for title, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞")
    run.font.size = Pt(11)
    add_footer_note(doc)
    doc.save(OUT_DIR / "nensho_kojin.docx")


def gen_nensho_xlsx():
    """念書 Excel（複数件管理）"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()

    # 一覧シート
    ws = wb.active
    ws.title = "念書一覧"
    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 18
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 14

    ws['B2'] = "念書 管理一覧"
    ws['B2'].font = Font(size=18, bold=True)

    ws['B4'] = "No."; ws['C4'] = "作成日"; ws['D4'] = "相手方氏名"; ws['E4'] = "金額"; ws['F4'] = "履行期日"; ws['G4'] = "ステータス"
    header_fill = PatternFill('solid', fgColor='E8F0FE')
    for col in ['B4','C4','D4','E4','F4','G4']:
        ws[col].font = Font(bold=True)
        ws[col].fill = header_fill
        ws[col].alignment = Alignment(horizontal='center')

    for i in range(5, 25):
        ws[f'B{i}'] = i - 4
        ws[f'B{i}'].alignment = Alignment(horizontal='center')

    # 念書本体シート
    ws2 = wb.create_sheet("念書本体")
    ws2.column_dimensions['A'].width = 4
    ws2.column_dimensions['B'].width = 30
    ws2.column_dimensions['C'].width = 50

    ws2['B2'] = "念　書"
    ws2['B2'].font = Font(size=22, bold=True)
    ws2['B2'].alignment = Alignment(horizontal='center')

    ws2['B4'] = "作成日"; ws2['C4'] = "令和　年　月　日"
    ws2['B5'] = "宛名"; ws2['C5'] = "○○　○○　殿"
    ws2['B7'] = "件名"; ws2['C7'] = "（金銭貸借/支払い/誓約/損害賠償 等）"

    ws2['B9'] = "債務の確認"; ws2['C9'] = "（金額・発生日・原契約等を具体的に）"
    ws2['B11'] = "履行内容"; ws2['C11'] = "（金額・行為内容を具体的に）"
    ws2['B13'] = "履行期日"; ws2['C13'] = "令和　年　月　日"
    ws2['B15'] = "履行方法"; ws2['C15'] = "甲指定の銀行口座への振込（手数料は乙負担）"
    ws2['B17'] = "違反時の措置"; ws2['C17'] = "民事訴訟・強制執行を含む法的措置を執ることに同意"
    ws2['B19'] = "公正証書化"; ws2['C19'] = "甲が必要と認めた場合は協力する"

    ws2['B22'] = "住所"; ws2['C22'] = ""
    ws2['B23'] = "氏名"; ws2['C23'] = ""

    wb.save(OUT_DIR / "nensho.xlsx")


def _gen_ininjo(filename, title, body_purpose, witness_clause):
    """委任状の共通生成"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(22); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    doc.add_paragraph("")

    # 委任者
    p = doc.add_paragraph()
    run = p.add_run("【委任者】"); run.bold = True; run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　㊞\n（必要に応じて：生年月日／法人代表者・役職）")
    run.font.size = Pt(10)
    doc.add_paragraph("")

    # 受任者
    p = doc.add_paragraph()
    run = p.add_run("【受任者】"); run.bold = True; run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("住所：　　　　　　　　　　　　　　\n氏名：　　　　　　　　　　　　　　　\n（委任者との関係：　　　　）")
    run.font.size = Pt(10)
    doc.add_paragraph("")

    # 委任事項
    p = doc.add_paragraph()
    run = p.add_run("私は、上記の者を代理人と定め、下記の事項に関する一切の権限を委任します。")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【委任事項】"); run.bold = True; run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run(body_purpose); run.font.size = Pt(10)
    doc.add_paragraph("")

    if witness_clause:
        p = doc.add_paragraph()
        run = p.add_run(witness_clause); run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / filename)


def gen_ininjo_sokai():
    purpose = (
        "1. 議決権の行使に関する一切の権限\n"
        "2. 動議の提出および賛否表明\n"
        "3. その他、本総会の議事運営に関連する一切の事項\n"
        "\n"
        "■ 議決権行使指示（必要な場合のみ）\n"
        "・第1号議案： 賛成 / 反対 / 受任者一任\n"
        "・第2号議案： 賛成 / 反対 / 受任者一任\n"
        "・第3号議案： 賛成 / 反対 / 受任者一任"
    )
    witness = (
        "本委任状は、令和　年　月　日開催の○○総会のためにのみ有効とし、"
        "他の用途には使用しないものとします。"
    )
    _gen_ininjo("ininjo_sokai.docx", "委 任 状（総会用）", purpose, witness)


def gen_ininjo_ginko():
    purpose = (
        "下記の手続きおよびこれに付随する一切の事項。\n"
        "（該当に丸を付けてください）\n"
        "・口座開設 / 口座解約 / 残高証明書の発行\n"
        "・通帳・キャッシュカードの再発行\n"
        "・名義変更 / 住所変更\n"
        "・相続手続き（被相続人： 　　　　／死亡日：　　年　月　日）\n"
        "・その他：　　　　　　　　　　\n"
        "\n"
        "■ 対象口座\n"
        "・銀行名： 　　　　銀行 　　　　支店\n"
        "・口座種別： 普通／当座／定期\n"
        "・口座番号： 　　　　　　　　"
    )
    witness = (
        "本委任状の使用にあたっては、委任者の印鑑証明書（発行日から3ヶ月以内）および"
        "本人確認書類のコピーを添付するものとします。"
    )
    _gen_ininjo("ininjo_ginko.docx", "委 任 状（銀行手続き用）", purpose, witness)


def gen_ininjo_general():
    purpose = (
        "下記の手続きおよびこれに付随する一切の事項。\n"
        "・（具体的な手続き内容を記載）\n"
        "・（提出書類の受領・押印を含む）\n"
        "・（必要に応じて費用の支払い）\n"
        "\n"
        "■ 委任の範囲\n"
        "本委任は、上記事項の完了をもって終了します。"
    )
    witness = (
        "本委任状は、令和　年　月　日までの間有効とし、必要に応じて本人確認書類のコピーを添付します。"
    )
    _gen_ininjo("ininjo_general.docx", "委 任 状", purpose, witness)


def gen_ryoshusho_oshare_xlsx():
    """領収書 おしゃれデザイン Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "領収書"
    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 18

    # タイトル（グラデーション風ヘッダー）
    ws['B2'] = "RECEIPT"
    ws['B2'].font = Font(size=22, bold=True, color='1A2940')
    ws.merge_cells('B2:D2')
    ws['B2'].alignment = Alignment(horizontal='center')

    ws['B3'] = "領　収　書"
    ws['B3'].font = Font(size=14, bold=True, color='C9A961')
    ws.merge_cells('B3:D3')
    ws['B3'].alignment = Alignment(horizontal='center')

    # 発行日・No.
    ws['B5'] = "No."; ws['C5'] = "0001"
    ws['B6'] = "発行日"; ws['C6'] = "令和　年　月　日"

    # 宛名
    ws['B8'] = "宛名"; ws['C8'] = "○○ 様"
    ws['C8'].font = Font(size=14, bold=True)

    # 金額
    ws['B10'] = "金額"; ws['C10'] = "￥　　　　　　　-"
    ws['C10'].font = Font(size=20, bold=True, color='1A2940')
    ws['B10'].fill = PatternFill('solid', fgColor='FAF8F4')

    # 但し書き
    ws['B12'] = "但し"; ws['C12'] = "（商品/サービス内容を記載）として"
    ws['B13'] = ""; ws['C13'] = "上記正に領収いたしました。"

    # インボイス情報
    ws['B15'] = "登録番号"; ws['C15'] = "T1234567890123"
    ws['B16'] = "内訳"; ws['C16'] = "10%対象 ￥　　　　／ 8%対象 ￥　　　　"

    # 発行者
    ws['B19'] = "発行者"; ws['C19'] = "○○商店"
    ws['B20'] = ""; ws['C20'] = "東京都○○区○○ 1-2-3"
    ws['B21'] = ""; ws['C21'] = "代表者：○○ ○○"

    # ノート
    ws['B23'] = "本テンプレートは5デザインパターン同梱（カラーロゴ挿入対応）"
    ws['B23'].font = Font(color='808080', italic=True)

    wb.save(OUT_DIR / "ryoshusho_oshare.xlsx")


def gen_ryoshusho_a4_4mai_xlsx():
    """領収書 A4 4枚割付 Excel（連番自動）"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "A4 4枚割付"
    for col in ['A','B','C','D','E','F','G','H']:
        ws.column_dimensions[col].width = 14

    thin = Side(border_style="thin", color="000000")
    box = Border(top=thin, bottom=thin, left=thin, right=thin)

    # タイトル
    ws['A1'] = "領収書（A4 4枚割付・連番自動・小売向け）"
    ws['A1'].font = Font(size=11, bold=True)

    # 4枚を 2x2 に配置（簡易版）
    # 1枚目: A3:D14
    # 2枚目: E3:H14
    # 3枚目: A16:D27
    # 4枚目: E16:H27

    def write_receipt(start_row, start_col, num):
        """指定位置に1枚分の領収書を書き込む"""
        col_offset = {1: 'A', 5: 'E'}[start_col]
        col_value = {1: 'B', 5: 'F'}[start_col]
        col_end = {1: 'D', 5: 'H'}[start_col]

        ws[f'{col_offset}{start_row}'] = "領 収 書"
        ws[f'{col_offset}{start_row}'].font = Font(size=14, bold=True)
        ws.merge_cells(f'{col_offset}{start_row}:{col_end}{start_row}')

        ws[f'{col_offset}{start_row+1}'] = "No."
        ws[f'{col_value}{start_row+1}'] = num
        ws[f'{col_offset}{start_row+2}'] = "発行日"
        ws[f'{col_value}{start_row+2}'] = "/  /  "
        ws[f'{col_offset}{start_row+3}'] = "宛名"
        ws[f'{col_value}{start_row+3}'] = "様"
        ws[f'{col_offset}{start_row+4}'] = "金額"
        ws[f'{col_value}{start_row+4}'] = "￥"
        ws[f'{col_offset}{start_row+5}'] = "但し"
        ws[f'{col_value}{start_row+5}'] = "として"
        ws[f'{col_offset}{start_row+6}'] = "登録番号"
        ws[f'{col_value}{start_row+6}'] = "T"
        ws[f'{col_offset}{start_row+7}'] = "発行者"
        ws[f'{col_value}{start_row+7}'] = ""
        ws[f'{col_offset}{start_row+9}'] = "- - - - 切り取り線 - - - -"
        ws[f'{col_offset}{start_row+9}'].font = Font(size=9, italic=True, color='808080')

    write_receipt(3, 1, "0001")
    write_receipt(3, 5, "0002")
    write_receipt(16, 1, "0003")
    write_receipt(16, 5, "0004")

    # 連番自動（管理シート）
    ws2 = wb.create_sheet("連番管理")
    ws2.column_dimensions['A'].width = 8
    ws2.column_dimensions['B'].width = 30
    ws2['A1'] = "連番"; ws2['B1'] = "発行日"
    ws2['A1'].font = Font(bold=True); ws2['B1'].font = Font(bold=True)
    for i in range(2, 21):
        ws2[f'A{i}'] = f'=TEXT({i-1},"0000")'

    wb.save(OUT_DIR / "ryoshusho_a4_4mai.xlsx")


def gen_ryoshusho_word():
    """領収書 Word（差し込み印刷対応）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("領 収 書")
    run.font.size = Pt(22); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("No.　0001\n発行日：令和　年　月　日")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("【宛名】　《差込: 取引先名》　様")
    run.font.size = Pt(13); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("金額：￥　《差込: 金額》　-")
    run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")

    sections = [
        "但し　《差込: 但し書き内容》として",
        "上記正に領収いたしました。",
        "",
        "登録番号： T1234567890123（適格請求書発行事業者）",
        "内訳： 10%対象 ￥（自動計算）／ 8%対象 ￥（自動計算）",
        "",
        "■ 差し込み印刷の使い方",
        "1. Word の「差し込み文書」タブ → 「宛先の選択」",
        "2. Excel データソース（取引先名・金額・但し書き）を指定",
        "3. 「結果のプレビュー」で確認",
        "4. 「完了と差し込み」 → 「文書の印刷」",
    ]
    for line in sections:
        p = doc.add_paragraph()
        run = p.add_run(line); run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("発行者： ○○商店\n所在地： 東京都○○区○○ 1-2-3\n代表者： ○○ ○○ ㊞")
    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "ryoshusho.docx")


def gen_ryoshusho_soejou():
    """領収書 添え状（送付状）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n○○部　○○　○○　様")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("○○株式会社\n○○部　○○　○○\nTEL：　　-　　-　　")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("領収書送付のご案内")
    run.font.size = Pt(16); run.bold = True
    doc.add_paragraph("")

    body = (
        "拝啓　時下ますますご清栄のこととお慶び申し上げます。\n"
        "平素は格別のご高配を賜り、厚く御礼申し上げます。\n\n"
        "さて、ご請求金額につきまして、お振込みを確認させていただきました。\n"
        "つきましては、下記のとおり領収書を同封させていただきましたので、ご査収のほどよろしくお願い申し上げます。\n\n"
        "今後とも変わらぬお引き立てを賜りますようお願い申し上げます。\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(body); run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("敬具"); run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("記"); run.bold = True; run.font.size = Pt(12)

    items = (
        "・領収書（No.　　　　　／ 金額：￥　　　　-）　　　　1通\n"
        "・その他添付書類（必要に応じて記載）　　　　　　　　　1通\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(items); run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("以上"); run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "ryoshusho_soejou.docx")


def gen_oreijo_english():
    """お礼状（英語版・Thank You Letter）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("May 5, 2026"); run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("Mr./Ms. ___________\nABC Corporation\nAddress: ___________________________"); run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Letter of Appreciation"); run.font.size = Pt(16); run.bold = True
    doc.add_paragraph("")

    body = (
        "Dear Mr./Ms. ___________,\n\n"
        "I am writing to express my sincere gratitude for your support of (project / meeting / occasion).\n"
        "Your professional advice and prompt cooperation have been invaluable to the successful completion of (specific work / contract).\n\n"
        "I look forward to continuing our productive working relationship in the future.\n\n"
        "Please do not hesitate to contact me if you have any questions or require further information.\n\n"
        "Sincerely yours,\n\n"
        "[Your Name]\n"
        "[Your Title], [Your Company]\n"
        "Email: ____________  /  Tel: ____________\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(body); run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "oreijo_english.docx")


def gen_oreijo_ochugen():
    """お礼状（お中元・お歳暮）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n○○部　○○　○○　様"); run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("○○株式会社\n○○部　○○　○○"); run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("お中元（お歳暮）御礼状"); run.font.size = Pt(16); run.bold = True
    doc.add_paragraph("")

    body = (
        "拝啓　盛夏（または歳末）の候、貴社ますますご清栄のこととお慶び申し上げます。\n"
        "平素は格別のご高配を賜り、厚く御礼申し上げます。\n\n"
        "さて、このたびはご丁寧にお中元（お歳暮）の品をお送りいただき、誠にありがとうございました。\n"
        "結構なお品物をお贈りいただき、社員一同心より感謝申し上げます。\n\n"
        "本来であれば直接お伺いしてお礼を申し上げるべきところ、書中をもちまして失礼ながらご挨拶とさせていただきます。\n\n"
        "今後とも変わらぬお引き立てを賜りますようお願い申し上げます。\n"
        "末筆ながら、貴社のますますのご発展と皆様のご健勝を心よりお祈り申し上げます。\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(body); run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("敬具"); run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "oreijo_ochugen.docx")


def gen_shanai_tsutatsu():
    """社内通達"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("通達番号： 　　-　　\n発行日： 令和　年　月　日"); run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("社 内 通 達"); run.font.size = Pt(20); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("各位\n（または： 部署別宛先・○○部 全員）"); run.font.size = Pt(11)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("発信者： ○○部　○○　○○"); run.font.size = Pt(11)
    doc.add_paragraph("")

    sections = [
        ("【件名】", "（例：人事異動のお知らせ／システム障害発生に伴うご連絡／福利厚生制度改定のご案内）"),
        ("【概要】", "本件の要点を3行以内で記載します。"),
        ("【適用日】", "令和　年　月　日　より"),
        ("【対象者】", "（部署・役職・全社員 等）"),
        ("【詳細】", "1. 変更点・実施事項を箇条書きで具体的に\n2. 関連するルール・規程の参照先\n3. 必要な手続きや対応期日"),
        ("【お問い合わせ】", "○○部　○○ ○○\n内線：　　　／メール：　　　@　　　"),
    ]
    for title, body in sections:
        p = doc.add_paragraph()
        run = p.add_run(title); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("\n本件について不明点があれば、上記担当者までお問い合わせください。\n本通達は ○○部 全員 への周知をお願いします。")
    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "shanai_tsutatsu.docx")


def gen_zaiko_kanri_xlsx():
    """在庫管理（スプレッドシート向け）Excel"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook()

    # 在庫一覧
    ws = wb.active
    ws.title = "在庫一覧"
    widths = {'A':4,'B':14,'C':24,'D':10,'E':10,'F':10,'G':14,'H':14}
    for col, w in widths.items(): ws.column_dimensions[col].width = w

    ws['B2'] = "在庫管理 一覧（自動連動）"
    ws['B2'].font = Font(size=18, bold=True)
    ws.merge_cells('B2:H2')

    headers = ["No.","SKU","商品名","在庫数","最小在庫","単価","合計金額","アラート"]
    for i, h in enumerate(headers):
        col = chr(ord('A') + i + 1)  # B,C,...
        cell = f'{col}4'
        ws[cell] = h
        ws[cell].font = Font(bold=True)
        ws[cell].fill = PatternFill('solid', fgColor='E8F0FE')
        ws[cell].alignment = Alignment(horizontal='center')

    # サンプルデータ
    sample = [
        (1, "A-001", "商品A", 50, 20, 1000, "=E5*G5", '=IF(E5<=F5,"発注!","")'),
        (2, "A-002", "商品B", 15, 20, 2000, "=E6*G6", '=IF(E6<=F6,"発注!","")'),
        (3, "B-001", "商品C", 100, 30, 500, "=E7*G7", '=IF(E7<=F7,"発注!","")'),
    ]
    for i, row in enumerate(sample, start=5):
        for j, val in enumerate(row):
            col = chr(ord('A') + j + 1)
            ws[f'{col}{i}'] = val

    # 入出庫履歴
    ws2 = wb.create_sheet("入出庫履歴")
    for col, w in {'A':4,'B':14,'C':14,'D':14,'E':10,'F':14,'G':30}.items():
        ws2.column_dimensions[col].width = w
    ws2['B2'] = "入出庫履歴"
    ws2['B2'].font = Font(size=14, bold=True)

    h2 = ["日付","区分","SKU","数量","担当者","備考"]
    for i, h in enumerate(h2):
        col = chr(ord('A') + i + 1)
        ws2[f'{col}4'] = h
        ws2[f'{col}4'].font = Font(bold=True)

    # 月次棚卸シート
    ws3 = wb.create_sheet("月次棚卸")
    ws3.column_dimensions['A'].width = 4
    ws3.column_dimensions['B'].width = 14
    ws3.column_dimensions['C'].width = 24
    ws3.column_dimensions['D'].width = 12
    ws3.column_dimensions['E'].width = 12
    ws3.column_dimensions['F'].width = 12
    ws3.column_dimensions['G'].width = 14
    ws3['B2'] = "月次棚卸シート"
    ws3['B2'].font = Font(size=14, bold=True)
    ws3['B4'] = "棚卸日"; ws3['C4'] = "令和　年　月　日"
    ws3['B6'] = "SKU"; ws3['C6'] = "商品名"; ws3['D6'] = "システム残"; ws3['E6'] = "実数"; ws3['F6'] = "差異"; ws3['G6'] = "備考"
    for h in ['B6','C6','D6','E6','F6','G6']:
        ws3[h].font = Font(bold=True)
        ws3[h].fill = PatternFill('solid', fgColor='E8F0FE')

    wb.save(OUT_DIR / "zaiko_kanri.xlsx")


# ==========================================================
# Vol.3: 法律深堀テンプレート7本
# ==========================================================
def _gen_houritsu_doc(filename, title, addressee_lines, sender_lines, clauses, footer_text=None):
    """Vol.3 法律書類の共通生成"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title); run.font.size = Pt(22); run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日"); run.font.size = Pt(11)
    doc.add_paragraph("")

    if addressee_lines:
        for line in addressee_lines:
            p = doc.add_paragraph()
            run = p.add_run(line); run.font.size = Pt(11)
        doc.add_paragraph("")

    for title_, body in clauses:
        p = doc.add_paragraph()
        run = p.add_run(title_); run.bold = True; run.font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run(body); run.font.size = Pt(10)

    if footer_text:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run(footer_text); run.font.size = Pt(10)

    if sender_lines:
        doc.add_paragraph("")
        for line in sender_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line); run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / filename)


def gen_naiyo_shomei():
    addressee = ["○○　○○　殿", "（住所：　　　　　　　　　　　　）"]
    sender = ["差出人 住所：　　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("件名", "（債権回収／損害賠償／契約解除／労働問題／近隣トラブル／賃貸トラブル 等を具体的に）"),
        ("本文", "私は、貴殿に対し、下記事項を通知いたします。\n\n（具体的な事実関係・請求内容・期日を1ページ520字以内で記載）\n\n本書面到達後○日以内に、貴殿の誠意ある対応を要求します。\n対応がない場合は、法的措置（民事訴訟・刑事告発・強制執行）を執らざるを得ないことを申し添えます。"),
        ("根拠条項", "（民法・商法・労働基準法・消費者契約法 等の根拠を明示）"),
        ("送付先", "本件に関するご連絡は、上記差出人までお願い申し上げます。"),
    ]
    _gen_houritsu_doc("naiyo_shomei_general.docx", "通　知　書", addressee, sender, clauses,
                     footer_text="※ 本書面は内容証明郵便（謄本3通・複写式）にて送付するものとします。")


def gen_shakuyosho():
    addressee = ["貸主：○○　○○　殿"]
    sender = ["借主 住所：　　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞", "（必要に応じて：連帯保証人）"]
    clauses = [
        ("第1条（金銭の借用）", "私は、本日、貴殿より下記金員を確かに借用いたしました。\n金　　　　　円也（金額は漢数字または改ざん防止のため算用数字＋全角スペースで明記）"),
        ("第2条（返済期日・方法）", "・返済期日：令和　年　月　日（一括 または 分割）\n・返済方法：貴殿指定の銀行口座への振込（振込手数料は借主負担）"),
        ("第3条（利息）", "利息は年　％（年率18%以下・利息制限法準拠）とし、元金と併せて返済する。"),
        ("第4条（遅延損害金）", "返済期日に遅延した場合、年14.6％の遅延損害金を支払うものとする。"),
        ("第5条（期限の利益喪失）", "1回でも支払いを怠った場合、残債務全額について期限の利益を失い、直ちに一括して支払う。"),
        ("第6条（連帯保証）", "（任意）連帯保証人は本借用書記載の債務について借主と連帯して責任を負う。"),
        ("第7条（公正証書化）", "貴殿が必要と認めた場合は、本借用書を公正証書として作成することに協力する。"),
    ]
    _gen_houritsu_doc("shakuyosho_general.docx", "借　用　書", addressee, sender, clauses,
                     footer_text="※ 借用金額1万円以上の場合、印紙税法に基づく収入印紙の貼付が必要です。")


def gen_kaiko_tsuchisho():
    addressee = ["○○　○○　殿"]
    sender = ["○○株式会社", "代表取締役　○○　○○", "印"]
    clauses = [
        ("件名", "解雇通知書"),
        ("通知事項", "貴殿との雇用契約について、下記の通り解雇いたしますので通知いたします。"),
        ("解雇日", "令和　年　月　日　（解雇予告日から30日以上経過後・労働基準法第20条準拠）"),
        ("解雇の種類", "□ 普通解雇 　□ 整理解雇 　□ 懲戒解雇 　□ その他（　　　　）"),
        ("解雇理由", "（具体的事由を明記。労働基準法第22条に基づき、本人の請求があれば解雇理由証明書を遅滞なく交付）"),
        ("解雇予告手当", "労働基準法第20条に基づき、解雇予告期間が30日に満たない場合は不足日数分の平均賃金を支払う。"),
        ("退職金等", "退職金、未払賃金、有給休暇買取等は別途精算書にて精算します。"),
        ("最終出勤日", "令和　年　月　日"),
        ("貸与物の返還", "社員証・健康保険証・PC・名刺・制服等を最終出勤日までに返還してください。"),
    ]
    _gen_houritsu_doc("kaiko_tsuchisho_general.docx", "解 雇 通 知 書", addressee, sender, clauses,
                     footer_text="※ 解雇は労働者にとって重大な不利益処分です。労働契約法第16条（解雇権濫用法理）に基づき、客観的合理性・社会通念上の相当性が必要です。")


def gen_taishoku_shomeisho():
    addressee = ["○○　○○　殿"]
    sender = ["○○株式会社", "代表取締役　○○　○○", "印"]
    clauses = [
        ("件名", "退職証明書"),
        ("証明事項（労働基準法第22条準拠）", "下記事項について、貴殿が在職していたことを証明します。"),
        ("1. 在職期間", "令和　年　月　日　〜　令和　年　月　日"),
        ("2. 業務の種類", "（職種・部署・役職）"),
        ("3. その事業における地位", "（一般従業員／管理職／役員 等）"),
        ("4. 賃金", "（月給・年俸・基本給／必要に応じて記載）"),
        ("5. 退職の事由", "□ 自己都合退職 　□ 会社都合退職 　□ 解雇 　□ 定年 　□ その他（　　　　）"),
    ]
    _gen_houritsu_doc("taishoku_shomeisho_general.docx", "退 職 証 明 書", addressee, sender, clauses,
                     footer_text="※ 労働基準法第22条第1項により、労働者の請求があった場合に遅滞なく交付する義務があります。労働者が請求しない事項は記載できません（同条第3項）。")


def gen_jidan_sho():
    addressee = []
    sender = ["甲（被害者）住所：　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞", "", "乙（加害者）住所：　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("（前文）", "甲（被害者）と乙（加害者）は、令和　年　月　日に発生した事故（または紛争）について、下記のとおり示談いたします。"),
        ("第1条（事故・紛争の内容）", "（発生日時・場所・態様を具体的に記載）"),
        ("第2条（示談金額）", "乙は甲に対し、本件示談金として金　　　　　円を支払う。"),
        ("第3条（支払い方法）", "・支払い期日：令和　年　月　日\n・支払い方法：甲指定の銀行口座への振込（手数料は乙負担）"),
        ("第4条（清算条項）", "甲乙双方は、本件に関し本示談書に定める他、何らの債権債務関係がないことを相互に確認する。"),
        ("第5条（守秘義務）", "甲乙双方は、本示談の存在および内容を、第三者（家族・専門家を除く）に開示しない。"),
        ("第6条（違反時の措置）", "本示談に違反した場合、違反した当事者は相手方に対し違約金　　　　円を支払う。"),
        ("第7条（公正証書化）", "甲乙協議のうえ、本示談書を公正証書として作成することに協力する。"),
    ]
    _gen_houritsu_doc("jidan_sho_general.docx", "示　談　書", addressee, sender, clauses)


def gen_yuigon_sho():
    addressee = []
    sender = ["遺言者 住所：　　　　　　　　　　　　", "生年月日：昭和（平成）　年　月　日", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("（前文）", "遺言者○○○○は、本日、次のとおり遺言する。"),
        ("第1条（不動産の遺贈）", "遺言者は、下記不動産を○○○○（相続人または受遺者の氏名）に相続させる（または遺贈する）。\n・所在：　　　　　　　　　　　　\n・地番：　　　　　／地目：　　　　　／地積：　　　　㎡"),
        ("第2条（預貯金の遺贈）", "遺言者は、下記預貯金を○○○○に相続させる。\n・○○銀行 ○○支店 普通預金 口座番号 　　　　　　"),
        ("第3条（有価証券・動産）", "遺言者は、（株式・投資信託・自動車・骨董品 等）を○○○○に相続させる。"),
        ("第4条（遺言執行者の指定）", "遺言者は、本遺言の執行者として ○○○○（住所：　　　　／氏名：　　　　　）を指定する。\n遺言執行者は、本遺言を執行するために必要な一切の権限を有する。"),
        ("第5条（付言事項）", "（家族への感謝の言葉・遺言の趣旨など、自由に記載）"),
    ]
    _gen_houritsu_doc("yuigon_sho.docx", "遺　言　書", addressee, sender, clauses,
                     footer_text="※ 自筆証書遺言は、遺言者が全文・日付・氏名を自書し、これに押印しなければ無効となる（民法第968条）。財産目録のみパソコン作成可（同条第2項）。法務局保管制度（2020年7月施行）の活用を推奨。")


def gen_rikon_kyogisho():
    addressee = []
    sender = ["甲（夫）住所：　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞", "", "乙（妻）住所：　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("（前文）", "甲（夫）○○○○と乙（妻）○○○○は、本日、協議離婚するにあたり、下記の通り合意した。"),
        ("第1条（離婚）", "甲乙は、本日、協議離婚する。離婚届の提出は乙が行い、甲はこれに協力する。"),
        ("第2条（親権）", "未成年の子（氏名：　　　／生年月日：　　年　月　日）の親権者は、（甲／乙）と定める。"),
        ("第3条（養育費）", "甲は乙に対し、子の養育費として、令和　年　月から子が満20歳に達する日の属する月まで、毎月末日限り、金　　　　円を、乙指定の銀行口座に振込により支払う。"),
        ("第4条（面会交流）", "甲は、子と月　回程度の頻度で面会交流することができる。具体的な日時・場所・方法は、甲乙協議のうえ決定する。"),
        ("第5条（財産分与）", "甲乙双方の婚姻中に形成した財産について、（具体的な分与内容を記載：不動産・預貯金・自動車・退職金・年金分割など）"),
        ("第6条（慰謝料）", "（必要に応じて）甲は乙に対し、慰謝料として金　　　　円を令和　年　月　日までに、乙指定の銀行口座に振込により支払う。"),
        ("第7条（年金分割）", "甲乙は、厚生年金保険法に基づく合意分割（按分割合：　　）を行う。"),
        ("第8条（清算条項）", "甲乙双方は、本協議書に定める他、互いに財産上・身分上の請求をしないことを相互に確認する。"),
        ("第9条（強制執行認諾文言）", "本協議書を公正証書とする場合、養育費・慰謝料の不履行があった場合は直ちに強制執行に服することを認諾する。"),
    ]
    _gen_houritsu_doc("rikon_kyogisho.docx", "離 婚 協 議 書", addressee, sender, clauses,
                     footer_text="※ 養育費・慰謝料の確実な履行のため、本協議書は公正証書化することを強く推奨します（強制執行認諾文言の付与により、不履行時に裁判なしで強制執行が可能）。")


def gen_naiyo_shomei_saiken():
    addressee = ["○○　○○　殿", "（住所：　　　　　　　　　　　　）"]
    sender = ["差出人 住所：　　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("件名", "債権回収のお願い（内容証明郵便）"),
        ("債権の特定", "私は、貴殿に対し、下記債権を有しております。\n・原契約：令和　年　月　日付け（売買契約／金銭消費貸借／業務委託 等）\n・債権額：金　　　　　円\n・支払期日：令和　年　月　日（既に経過）"),
        ("請求事項", "本書面到達後 7 日以内に、上記金員を私の指定する銀行口座に一括にてお支払い下さい。"),
        ("法的措置の予告", "上記期間内にお支払いがない場合、誠に遺憾ながら下記法的措置を執らせていただきます。\n（1）民事訴訟（少額訴訟・通常訴訟）\n（2）強制執行（給与差押え・預金差押え・動産差押え）\n（3）必要に応じて刑事告訴（詐欺罪等）"),
        ("時効更新", "本書面の到達をもって、民法第150条第1項の催告として、消滅時効の完成を6ヶ月猶予するものとします。"),
    ]
    _gen_houritsu_doc("naiyo_shomei_saiken.docx", "通　知　書（債権回収）", addressee, sender, clauses,
                     footer_text="※ 本書面は内容証明郵便（謄本3通・複写式）にて送付するものとします。")


def gen_shakuyosho_kazoku():
    addressee = ["貸主：○○　○○（続柄：父／母／兄弟姉妹 等）　殿"]
    sender = ["借主 住所：　　　　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("第1条（金銭の借用）", "私は、本日、家族である貸主より下記金員を確かに借用いたしました。\n金　　　　　円也"),
        ("第2条（贈与税対策・最重要）", "本借用は贈与ではなく金銭消費貸借契約であり、後日、税務署等の指摘があった場合に贈与税課税を回避するため、下記事項を明記する。"),
        ("第3条（返済期日・方法）", "・返済期日：令和　年　月　日\n・返済方法：貸主指定の銀行口座への振込（振込記録を残すこと）\n・返済原資：借主の給与・事業収入等"),
        ("第4条（利息）", "金銭の貸借に伴う利息は年　％（無利息でも可・ただし家族間でも通常は年0.1〜2%程度を設定し贈与税認定リスクを下げる）"),
        ("第5条（遅延損害金）", "返済期日に遅延した場合、年14.6％の遅延損害金を支払う。"),
        ("第6条（連帯保証・任意）", "（家族間では通常不要だが、必要に応じて連帯保証人を立てる）"),
        ("第7条（贈与税回避のための重要事項）", "本借用書は、税務調査時の証拠資料となるため、必ず以下を実施する。\n（1）借入時に銀行振込で受領（現金手渡し不可）\n（2）毎月の返済を銀行振込で実行し記録を残す\n（3）借入時の収入印紙貼付（1万円以上）"),
    ]
    _gen_houritsu_doc("shakuyosho_kazoku.docx", "借用書（家族間）", addressee, sender, clauses,
                     footer_text="※ 家族間の金銭貸借は税務署から「贈与」と認定されやすいため、本借用書を必ず作成し、銀行振込による返済記録を残してください。")


def gen_kaiko_yokoku():
    addressee = ["○○　○○　殿"]
    sender = ["○○株式会社", "代表取締役　○○　○○", "印"]
    clauses = [
        ("件名", "解雇予告通知書"),
        ("通知事項", "貴殿との雇用契約について、下記の通り解雇予告いたしますので通知いたします。"),
        ("予告日", "令和　年　月　日（本通知書の到達日）"),
        ("解雇予定日", "令和　年　月　日（予告日から30日後・労働基準法第20条準拠）"),
        ("解雇の種類・理由", "□ 普通解雇 　□ 整理解雇 　□ 懲戒解雇\n（具体的事由を明記）"),
        ("予告期間中の処遇", "予告日から解雇予定日までの期間は、通常通り勤務していただきます。"),
        ("解雇理由証明書", "労働基準法第22条に基づき、本人の請求があれば解雇理由証明書を遅滞なく交付します。"),
        ("最終出勤日・引継ぎ", "令和　年　月　日。引継ぎを完了の上、貸与物（社員証・健康保険証・PC・名刺・制服等）を返還してください。"),
    ]
    _gen_houritsu_doc("kaiko_tsuchisho_yokoku.docx", "解 雇 予 告 通 知 書", addressee, sender, clauses,
                     footer_text="※ 労働基準法第20条により、解雇予告は30日以上前に行う必要があります。30日に満たない場合は、不足日数分の平均賃金を解雇予告手当として支払う必要があります。")


def gen_jidan_sho_jiko():
    addressee = []
    sender = ["甲（被害者）住所：　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞", "", "乙（加害者）住所：　　　　　　　　　", "氏名：　　　　　　　　　　　　　　　㊞"]
    clauses = [
        ("（前文）", "甲（被害者）と乙（加害者）は、令和　年　月　日に発生した交通事故について、下記のとおり示談する。"),
        ("第1条（事故の特定）", "・発生日時：令和　年　月　日　午前／午後　時　分頃\n・発生場所：　　　　　　　　　　　　\n・事故車両：甲（車種／ナンバー）／ 乙（車種／ナンバー）\n・事故態様：（追突／出会い頭／右直 等を具体的に）"),
        ("第2条（過失割合）", "甲：乙＝　：　（協議の結果合意した過失割合）"),
        ("第3条（人的損害）", "甲の人的損害（治療費・通院費・休業損害・慰謝料・後遺障害逸失利益）として、乙は甲に対し金　　　　　円を支払う。"),
        ("第4条（物的損害）", "甲の物的損害（車両修理費・代車費・買替差額）として、乙は甲に対し金　　　　　円を支払う。"),
        ("第5条（支払い方法）", "・支払い期日：令和　年　月　日\n・支払い方法：甲指定の銀行口座への振込（手数料は乙負担）\n・自賠責保険・任意保険からの支払いを充当する場合の調整方法"),
        ("第6条（清算条項）", "甲乙双方は、本件交通事故に関し本示談書に定める他、何らの債権債務関係がないことを相互に確認する。後遺障害が新たに発覚した場合の取扱いは別途協議する。"),
        ("第7条（守秘義務・違約金）", "甲乙双方は、本示談の存在および内容を、第三者に開示しない。違反した場合は違約金　　　　円を支払う。"),
    ]
    _gen_houritsu_doc("jidan_sho_jiko.docx", "示 談 書（交通事故）", addressee, sender, clauses,
                     footer_text="※ 交通事故示談は、過失割合・損害額の算定が複雑です。後遺障害発覚リスクや時効（人身5年・物損3年）を考慮し、必ず弁護士相談の上、示談書を作成してください。")


def gen_oreijo_general():
    """お礼状（汎用ビジネス）"""
    doc = Document()
    set_a4_portrait(doc.sections[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("○○株式会社\n○○部　○○　○○　様")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("○○株式会社\n○○部　○○　○○")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("お礼状")
    run.font.size = Pt(18); run.bold = True
    doc.add_paragraph("")

    body = (
        "拝啓　時下ますますご清栄のこととお慶び申し上げます。\n"
        "平素は格別のご高配を賜り、厚く御礼申し上げます。\n\n"
        "さて、このたびは（用件・受領内容を具体的に）にあたり、"
        "格別のご配慮を賜り誠にありがとうございました。"
        "おかげさまをもちまして滞りなく完了することができ、心より感謝申し上げます。\n\n"
        "つきましては、書面にて略儀ながら御礼申し上げます。\n"
        "今後とも変わらぬお引き立てを賜りますようお願い申し上げます。\n\n"
        "末筆ながら、貴社のますますのご発展と皆様のご健勝を心よりお祈り申し上げます。\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(body)
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("敬具")
    run.font.size = Pt(11)

    add_footer_note(doc)
    doc.save(OUT_DIR / "oreijo_general.docx")


def gen_mitsumori_iraisho_word():
    """見積依頼書 Word版"""
    doc = Document()
    set_a4_portrait(doc.sections[0])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("見　積　依　頼　書")
    run.font.size = Pt(22)
    run.bold = True
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("令和　年　月　日")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("　　　　　　　御中")
    run.font.size = Pt(11)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("拝啓　時下ますますご清栄のこととお慶び申し上げます。\n下記の通り見積をご依頼申し上げます。よろしくお願い申し上げます。")
    run.font.size = Pt(10)
    doc.add_paragraph("")

    items = [
        ("件名", ""),
        ("依頼内容", ""),
        ("仕様・条件", ""),
        ("数量", ""),
        ("納期", "令和　年　月　日"),
        ("納入場所", ""),
        ("見積回答期限", "令和　年　月　日"),
        ("支払条件", ""),
        ("その他", ""),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(items):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("【依頼者】\n会社名：\n所属：\n担当者：\n電話：\nメール：")
    run.font.size = Pt(10)

    add_footer_note(doc)
    doc.save(OUT_DIR / "mitsumori-iraisho.docx")


def gen_mitsumori_iraisho_xlsx():
    """見積依頼書 Excel版"""
    wb = Workbook()
    ws = wb.active
    ws.title = "見積依頼書"
    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill('solid', fgColor='F0F0F0')

    ws['B2'] = "見積依頼書"
    ws['B2'].font = Font(size=18, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('B2:E2')

    ws['D4'] = "依頼日："; ws['D4'].font = Font(bold=True)
    ws['E4'] = "令和　年　月　日"
    ws['B5'] = "宛先："; ws['B5'].font = Font(bold=True)
    ws['C5'] = "　　　　　御中"

    items = [
        ("件名", ""),
        ("依頼内容", ""),
        ("仕様・条件", ""),
        ("数量", ""),
        ("納期", "令和　年　月　日"),
        ("納入場所", ""),
        ("見積回答期限", "令和　年　月　日"),
        ("支払条件", ""),
        ("その他", ""),
    ]
    start_row = 7
    ws.cell(start_row, 2, "項目").font = Font(bold=True)
    ws.cell(start_row, 2).fill = header_fill
    ws.cell(start_row, 2).border = border
    ws.cell(start_row, 3, "内容").font = Font(bold=True)
    ws.cell(start_row, 3).fill = header_fill
    ws.cell(start_row, 3).border = border
    ws.merge_cells(start_row=start_row, start_column=3, end_row=start_row, end_column=5)
    for i, (k, v) in enumerate(items):
        r = start_row + 1 + i
        ws.cell(r, 2, k).border = border
        ws.cell(r, 3, v).border = border
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=5)

    ws['B18'] = "依頼者"
    ws['B18'].font = Font(bold=True)
    ws['B19'] = "会社名："
    ws['B20'] = "所属："
    ws['B21'] = "担当者："
    ws['B22'] = "電話："
    ws['B23'] = "メール："

    ws.column_dimensions['B'].width = 22
    for col in ['C', 'D', 'E']: ws.column_dimensions[col].width = 18

    note = wb.create_sheet("使い方")
    note['A1'] = "見積依頼書（RFQ）使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. 宛先・依頼日を記入"
    note['A4'] = "2. 件名・依頼内容・仕様・数量・納期・納入場所を具体的に記入"
    note['A5'] = "3. 見積回答期限を設定（業者が見積を返答する期限）"
    note['A6'] = "4. 依頼者情報を記入し、業者にPDFまたは紙で送付"
    note['A8'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A8'].font = Font(color='808080')

    wb.save(OUT_DIR / "mitsumori-iraisho.xlsx")


def gen_koyou_keiyakusho_xlsx():
    """雇用契約書 Excel版"""
    wb = Workbook()
    ws = wb.active
    ws.title = "雇用契約書"
    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill('solid', fgColor='F0F0F0')
    ws['B2'] = "雇用契約書"
    ws['B2'].font = Font(size=18, bold=True)
    ws['B2'].alignment = Alignment(horizontal='center')
    ws.merge_cells('B2:E2')
    ws['B4'] = "甲（雇用主）："
    ws['B5'] = "乙（労働者）："
    items = [
        ("契約期間", ""),
        ("試用期間", ""),
        ("身分（正社員/契約/アルバイト/嘱託）", ""),
        ("就業場所", ""),
        ("業務内容", ""),
        ("始業時刻", ""),
        ("終業時刻", ""),
        ("休憩時間", ""),
        ("所定休日", ""),
        ("年次有給休暇", ""),
        ("基本給（月給）", ""),
        ("諸手当", ""),
        ("固定残業代", ""),
        ("賃金支払日", ""),
        ("社会保険", ""),
        ("退職事項", ""),
        ("競業避止義務", ""),
        ("秘密保持", ""),
    ]
    start_row = 8
    ws.cell(start_row, 2, "項目").font = Font(bold=True)
    ws.cell(start_row, 2).fill = header_fill
    ws.cell(start_row, 2).border = border
    ws.cell(start_row, 3, "内容").font = Font(bold=True)
    ws.cell(start_row, 3).fill = header_fill
    ws.cell(start_row, 3).border = border
    ws.merge_cells(start_row=start_row, start_column=3, end_row=start_row, end_column=5)
    for i, (k, v) in enumerate(items):
        r = start_row + 1 + i
        ws.cell(r, 2, k).border = border
        ws.cell(r, 3, v).border = border
        ws.merge_cells(start_row=r, start_column=3, end_row=r, end_column=5)
    ws.column_dimensions['B'].width = 32
    for col in ['C','D','E']: ws.column_dimensions[col].width = 16
    note = wb.create_sheet("使い方")
    note['A1'] = "雇用契約書（Excel版）使い方"
    note['A1'].font = Font(size=14, bold=True)
    note['A3'] = "1. C列の各項目に必要事項を入力してください"
    note['A4'] = "2. 「身分」を正社員/契約/アルバイト/嘱託 から選んで記入"
    note['A5'] = "3. 試用期間・固定残業代・競業避止義務は該当がなければ「なし」で記入"
    note['A7'] = "配布元: テンプレートフリー（template-free.jp）"
    note['A7'].font = Font(color='808080')
    wb.save(OUT_DIR / "koyou-keiyakusho.xlsx")


# ==========================================================
# Main
# ==========================================================
if __name__ == "__main__":
    print("Generating templates...")

    # T001
    gen_taishokutodoke_a4_tategaki()
    print("  [ok]taishokutodoke_a4_tategaki.docx")
    gen_taishokutodoke_a4_yokogaki()
    print("  [ok]taishokutodoke_a4_yokogaki.docx")
    gen_taishokutodoke_b5_tategaki()
    print("  [ok]taishokutodoke_b5_tategaki.docx")
    gen_taishokutodoke_xlsx()
    print("  [ok]taishokutodoke.xlsx")

    # T011
    gen_roudoujouken_seishain()
    print("  [ok]roudoujouken-tsuuchisho_seishain.docx")
    gen_roudoujouken_keiyaku()
    print("  [ok]roudoujouken-tsuuchisho_keiyaku.docx")
    gen_roudoujouken_arbeit()
    print("  [ok]roudoujouken-tsuuchisho_arbeit.docx")
    gen_roudoujouken_xlsx()
    print("  [ok]roudoujouken-tsuuchisho.xlsx")

    # T016
    gen_gyoumu_itaku_general()
    print("  [ok]gyoumu-itaku_general.docx")
    gen_gyoumu_itaku_junin()
    print("  [ok]gyoumu-itaku_junin.docx")
    gen_gyoumu_itaku_ukeoi()
    print("  [ok]gyoumu-itaku_ukeoi.docx")

    # T002
    gen_taishokutodoke_muryou_a4()
    print("  [ok]taishokutodoke-muryou_a4.docx")
    gen_taishokutodoke_muryou_a4_tategaki()
    print("  [ok]taishokutodoke-muryou_a4_tategaki.docx")

    # T007
    gen_shimatsusho_general()
    print("  [ok]shimatsusho_general.docx")
    gen_shimatsusho_chikoku()
    print("  [ok]shimatsusho_chikoku.docx")
    gen_shimatsusho_mudankin()
    print("  [ok]shimatsusho_mudankin.docx")
    gen_shimatsusho_xlsx()
    print("  [ok]shimatsusho.xlsx")

    # T012
    gen_koyou_keiyakusho_seishain()
    print("  [ok]koyou-keiyakusho_seishain.docx")
    gen_koyou_keiyakusho_keiyaku()
    print("  [ok]koyou-keiyakusho_keiyaku.docx")
    gen_koyou_keiyakusho_arbeit()
    print("  [ok]koyou-keiyakusho_arbeit.docx")
    gen_koyou_keiyakusho_shokutaku()
    print("  [ok]koyou-keiyakusho_shokutaku.docx")
    gen_koyou_keiyakusho_xlsx()
    print("  [ok]koyou-keiyakusho.xlsx")

    # T023
    gen_seikyu_kojin_xlsx()
    print("  [ok]seikyu-kojin-jigyonushi.xlsx")
    gen_seikyu_kojin_word()
    print("  [ok]seikyu-kojin-jigyonushi.docx")

    # T029
    gen_mitsumori_xlsx()
    print("  [ok]mitsumori-tourokufuyou.xlsx")
    gen_mitsumori_word()
    print("  [ok]mitsumori-tourokufuyou.docx")

    # T035
    gen_gijiroku_standard()
    print("  [ok]gijiroku_standard.docx")
    gen_gijiroku_1on1()
    print("  [ok]gijiroku_1on1.docx")
    gen_gijiroku_brest()
    print("  [ok]gijiroku_brest.docx")
    gen_gijiroku_xlsx()
    print("  [ok]gijiroku.xlsx")

    # T020
    gen_himitsuhoji_nyusha()
    print("  [ok]himitsuhoji-seiyakusho_nyusha.docx")
    gen_himitsuhoji_taisha()
    print("  [ok]himitsuhoji-seiyakusho_taisha.docx")

    # T028
    gen_mitsumori_iraisho_word()
    print("  [ok]mitsumori-iraisho.docx")
    gen_mitsumori_iraisho_xlsx()
    print("  [ok]mitsumori-iraisho.xlsx")

    # T019: NDA（双方向 / 片務）
    gen_nda_souhou()
    print("  [ok]nda_souhou.docx")
    gen_nda_henmu()
    print("  [ok]nda_henmu.docx")

    # Vol.2 サイクル1: T041 念書 / T054 日報 / T057 お礼状
    gen_nensho_general()
    print("  [ok]nensho_general.docx")
    gen_nippo_standard()
    print("  [ok]nippo_standard.docx")
    gen_nippo_xlsx()
    print("  [ok]nippo.xlsx")
    gen_oreijo_general()
    print("  [ok]oreijo_general.docx")

    # Vol.2 サイクル2: T042 支払い念書 / T043 金銭念書 / T060 顛末書
    gen_nensho_shiharai()
    print("  [ok]nensho_shiharai.docx")
    gen_nensho_kinsen()
    print("  [ok]nensho_kinsen.docx")
    gen_tenmatsusho()
    print("  [ok]tenmatsusho_general.docx")

    # Vol.2 サイクル3: T044 念書個人特化 / T046 念書Excel
    gen_nensho_kojin()
    print("  [ok]nensho_kojin.docx")
    gen_nensho_xlsx()
    print("  [ok]nensho.xlsx")

    # Vol.2 サイクル4: T047 委任状総会 / T048 委任状銀行 / T049 委任状汎用
    gen_ininjo_sokai()
    print("  [ok]ininjo_sokai.docx")
    gen_ininjo_ginko()
    print("  [ok]ininjo_ginko.docx")
    gen_ininjo_general()
    print("  [ok]ininjo_general.docx")

    # Vol.2 サイクル5: T050 領収書おしゃれ / T051 領収書A4 4枚 / T052 領収書Word
    gen_ryoshusho_oshare_xlsx()
    print("  [ok]ryoshusho_oshare.xlsx")
    gen_ryoshusho_a4_4mai_xlsx()
    print("  [ok]ryoshusho_a4_4mai.xlsx")
    gen_ryoshusho_word()
    print("  [ok]ryoshusho.docx")

    # Vol.2 サイクル6: T053 領収書添え状（T055/T056 は既存ファイル流用）
    gen_ryoshusho_soejou()
    print("  [ok]ryoshusho_soejou.docx")

    # Vol.2 サイクル7: T058/T059/T061/T062
    gen_oreijo_english()
    print("  [ok]oreijo_english.docx")
    gen_oreijo_ochugen()
    print("  [ok]oreijo_ochugen.docx")
    gen_shanai_tsutatsu()
    print("  [ok]shanai_tsutatsu.docx")
    gen_zaiko_kanri_xlsx()
    print("  [ok]zaiko_kanri.xlsx")

    # Vol.3: 法律系深堀7本（T063/T067/T074/T077/T080/T082/T083）
    gen_naiyo_shomei()
    print("  [ok]naiyo_shomei_general.docx")
    gen_naiyo_shomei_saiken()
    print("  [ok]naiyo_shomei_saiken.docx")
    gen_shakuyosho()
    print("  [ok]shakuyosho_general.docx")
    gen_shakuyosho_kazoku()
    print("  [ok]shakuyosho_kazoku.docx")
    gen_kaiko_tsuchisho()
    print("  [ok]kaiko_tsuchisho_general.docx")
    gen_kaiko_yokoku()
    print("  [ok]kaiko_tsuchisho_yokoku.docx")
    gen_taishoku_shomeisho()
    print("  [ok]taishoku_shomeisho_general.docx")
    gen_jidan_sho()
    print("  [ok]jidan_sho_general.docx")
    gen_jidan_sho_jiko()
    print("  [ok]jidan_sho_jiko.docx")
    gen_yuigon_sho()
    print("  [ok]yuigon_sho.docx")
    gen_rikon_kyogisho()
    print("  [ok]rikon_kyogisho.docx")

    print(f"\nAll templates generated to: {OUT_DIR}")
